import base64
import io
import logging
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from bs4 import BeautifulSoup, NavigableString

logger = logging.getLogger("Monjez_Word_Engine")

def generate_local_word(html_content, letterhead_b64):
    doc = docx.Document()
    
    # 1. ضبط الإعدادات الافتراضية للمستند (خط رسمي مريح)
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(14)
    
    # 2. إدراج الرأسية في أعلى الصفحة وتوسيطها
    if letterhead_b64:
        try:
            img_data = base64.b64decode(letterhead_b64)
            img_stream = io.BytesIO(img_data)
            doc.add_picture(img_stream, width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            logger.error(f"خطأ في دمج الرأسية: {e}")

    # 3. تحليل هيكل HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # دالة الزحف الشجري الذكية
    def parse_element(element):
        # التعامل مع النصوص الحرة
        if isinstance(element, NavigableString):
            text = str(element).strip()
            if text:
                p = doc.add_paragraph(text)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            return

        # معالجة العناوين
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            text = element.get_text(separator=" ", strip=True)
            if text:
                level = min(int(element.name[1]), 9)
                p = doc.add_heading(text, level=level)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            return

        # معالجة الجداول
        if element.name == 'table':
            rows = element.find_all('tr')
            if rows:
                cols = max((len(row.find_all(['td', 'th'])) for row in rows), default=0)
                if cols > 0:
                    table = doc.add_table(rows=len(rows), cols=cols)
                    table.style = 'Table Grid' # حدود الجدول
                    for i, row in enumerate(rows):
                        cells = row.find_all(['td', 'th'])
                        for j, cell in enumerate(cells):
                            if j < cols:
                                cell_obj = table.cell(i, j)
                                # دمج النصوص داخل الخلية بشكل مرتب
                                cell_obj.text = cell.get_text(separator=" ", strip=True)
                                for paragraph in cell_obj.paragraphs:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            return

        # معالجة القوائم
        if element.name in ['ul', 'ol']:
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(li.get_text(separator=" ", strip=True), style='List Bullet')
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            return

        # معالجة الفقرات والـ Divs بشكل ذكي
        if element.name in ['p', 'div']:
            block_tags = ['p', 'div', 'table', 'ul', 'ol', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']
            # هل هذا العنصر هو مجرد غلاف يحتوي على عناصر كبيرة داخله؟
            has_block_children = any(child.name in block_tags for child in element.children if getattr(child, 'name', None))
            
            if has_block_children:
                # إذا كان غلافاً (مثل الغلاف الرئيسي للمستند)، غُص للداخل ولا تجمعه
                for child in element.children:
                    parse_element(child)
            else:
                # إذا كان فقرة نهائية أو div يحتوي على نصوص وتوقيعات (بدون جداول بداخله)
                text = element.get_text(separator=" ", strip=True)
                if text:
                    p = doc.add_paragraph(text)
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            return

        # أي وسوم أخرى (مثل span, section, b, strong) نغوص بداخلها
        if hasattr(element, 'children'):
            for child in element.children:
                parse_element(child)

    # بدء الزحف من جسم المستند
    root = soup.body if soup.body else soup
    for child in root.children:
        parse_element(child)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()
