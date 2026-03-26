import os
import re
import json
import logging
import base64
import time
import io
import concurrent.futures
import subprocess
import tempfile
from flask import Flask, request, jsonify

# ══════════════════════════════════════════════════════════
# ✅ استدعاء مكتبات الوورد المطلوبة للحقن العميق للرأسية وضبط الهوامش
# ══════════════════════════════════════════════════════════
import docx
from docx.shared import Inches, Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger("Monjez_V10_Server")

app = Flask(__name__)

# ── Lazy Gemini ──
_client = None
_types = None
_init = False

def get_client():
    global _client, _types, _init
    if not _init:
        _init = True
        try:
            from google import genai as g
            from google.genai import types as t
            _types = t
            k = os.environ.get("GOOGLE_API_KEY")
            if k:
                _client = g.Client(api_key=k, http_options={"api_version": "v1beta"})
                logger.info("✅ Monjez V10 Server (Ready)")
        except Exception as e:
            logger.error(f"Init: {e}")
    return _client

def get_types():
    get_client()
    return _types

def call_gemini(model, contents, config, timeout):
    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as ex:
        f = ex.submit(get_client().models.generate_content, model=model, contents=contents, config=config)
        return f.result(timeout=timeout)

# 💡 دالة جديدة لاستخراج الاستهلاك الدقيق للتوكنز
def extract_tokens(resp):
    try:
        if hasattr(resp, 'usage_metadata') and resp.usage_metadata:
            return getattr(resp.usage_metadata, 'total_token_count', 0)
    except Exception as e:
        logger.error(f"Token extraction error: {e}")
    return 0

def clean_html_output(raw_text):
    raw = raw_text.strip()
    if raw.startswith("`" * 3):
        raw = re.sub(r"^`{3}(?:html|xml)?\n?", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\n?`{3}$", "", raw)
    div_match = re.search(r'<div[^>]*xmlns="http://www.w3.org/1999/xhtml"[^>]*>(.*?)</div>\s*</foreignObject>', raw, re.DOTALL)
    if div_match:
        raw = div_match.group(1)
    raw = re.sub(r'\s?contenteditable="[^"]*"', '', raw, flags=re.IGNORECASE)
    raw = re.sub(r'\s?contenteditable=\'[^\']*\'', '', raw, flags=re.IGNORECASE)
    raw = re.sub(r'\s?contenteditable', '', raw, flags=re.IGNORECASE)
    return raw.strip()

# ══════════════════════════════════════════════════════════
# 🛡️ حقنة الجداول (درع الخطوط المزدوجة والصفوف الوهمية)
# ══════════════════════════════════════════════════════════
def force_table_borders(html_text):
    # 0. إزالة أوسمة البنية التي يُنشئها Gemini أحياناً وتسبب صفاً وهمياً في LibreOffice
    html_text = re.sub(r'</?thead[^>]*>', '', html_text, flags=re.IGNORECASE)
    html_text = re.sub(r'</?tbody[^>]*>', '', html_text, flags=re.IGNORECASE)
    html_text = re.sub(r'</?tfoot[^>]*>', '', html_text, flags=re.IGNORECASE)
    html_text = re.sub(r'<colgroup[^>]*>.*?</colgroup>', '', html_text, flags=re.IGNORECASE | re.DOTALL)
    html_text = re.sub(r'<caption[^>]*>.*?</caption>', '', html_text, flags=re.IGNORECASE | re.DOTALL)
    
    # 1. إجبار الجدول على التنسيق النظيف المندمج لمنع الخطوط المزدوجة
    html_text = html_text.replace("<table", "<table border='1' cellpadding='4' cellspacing='0' style='border-collapse:collapse; border-spacing:0; width:100%; border: 1px solid black; margin: 10px 0;' ")
    html_text = html_text.replace("<th", "<th style='border: 1px solid black; padding: 4px; text-align: center; vertical-align: middle; color: black;' ")
    html_text = html_text.replace("<td", "<td style='border: 1px solid black; padding: 4px; vertical-align: middle;' ")
    
    # 2. درع التنظيف: مسح أي صفوف فارغة (Empty Rows) أنشأها الذكاء الاصطناعي وتسبب الخانة الفارغة
    html_text = re.sub(r'<tr>\s*(?:<t[hd][^>]*>\s*</t[hd]>\s*)+</tr>', '', html_text, flags=re.IGNORECASE)
    html_text = re.sub(r'<tr>\s*(?:<t[hd][^>]*>\s*&nbsp;\s*</t[hd]>\s*)+</tr>', '', html_text, flags=re.IGNORECASE)
    # مسح صفوف فارغة تحتوي فقط على مسافات أو أسطر فارغة داخل الخلايا
    html_text = re.sub(r'<tr>\s*(?:<t[hd][^>]*>\s*(?:&nbsp;|\s)*</t[hd]>\s*)+</tr>', '', html_text, flags=re.IGNORECASE)
    
    return html_text

# ══════════════════════════════════════════════════════════
# 🔧 تحويل اتجاه الجداول إلى LTR قبل تصدير الوورد
#    (لأن bidiVisual في DOCX يتولى عكسها للعربية تلقائياً)
# ══════════════════════════════════════════════════════════
def force_tables_ltr_for_export(html_text):
    html_text = re.sub(r'(<table[^>]*?)\bdir\s*=\s*["\']rtl["\']', r'\1dir="ltr"', html_text, flags=re.IGNORECASE)
    return html_text

# ══════════════════════════════════════════════════════════
# 🚀 Local LibreOffice Converter
# ══════════════════════════════════════════════════════════
def local_libreoffice_convert(file_bytes, input_ext, output_ext):
    logger.info(f"🖥️ Local LibreOffice: Converting {input_ext.upper()} to {output_ext.upper()}...")
    
    # 🧹 [التعديل الجراحي]: تنظيف ملف قفل X99 لتجنب خطأ Server is already active for display 99
    lock_file = "/tmp/.X99-lock"
    if os.path.exists(lock_file):
        try:
            os.remove(lock_file)
            logger.info("🧹 Cleared stale /tmp/.X99-lock file.")
        except Exception as e:
            logger.warning(f"⚠️ Failed to clear lock file: {e}")

    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = os.path.join(temp_dir, f"input.{input_ext}")
            with open(input_path, 'wb') as f:
                f.write(file_bytes)
            
            profile_dir = os.path.join(temp_dir, "lo_profile")
            
            filters = {
                "docx": "docx:MS Word 2007 XML",
                "xlsx": "xlsx:Calc MS Excel 2007 XML",
                "pptx": "pptx:Impress MS PowerPoint 2007 XML",
                "pdf": "pdf"
            }
            lo_filter = filters.get(output_ext, output_ext)
            
            command = [
                'libreoffice',
                f'-env:UserInstallation=file://{profile_dir}',
                '--headless',
                '--invisible',
                '--nocrashreport',
                '--nodefault',
                '--nofirststartwizard',
                '--nologo',
                '--norestore',
                '--convert-to', lo_filter,
                '--outdir', temp_dir,
                input_path
            ]
            
            env = os.environ.copy()
            env["SAL_USE_VCLPLUGIN"] = "gen"
            env["LC_ALL"] = "C.UTF-8"
            env["DISPLAY"] = ":99"
            
            process = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=120, env=env)
            output_path = os.path.join(temp_dir, f"input.{output_ext}")
            
            if process.returncode == 0 and os.path.exists(output_path):
                with open(output_path, 'rb') as f:
                    result_bytes = f.read()
                logger.info("✅ Local LibreOffice: Conversion successful!")
                return result_bytes, None
            else:
                error_msg = process.stderr.decode('utf-8', errors='ignore').strip()
                if not error_msg:
                    error_msg = process.stdout.decode('utf-8', errors='ignore').strip() or "Unknown error"
                logger.error(f"❌ LibreOffice Failed! Code: {process.returncode}")
                logger.error(f"❌ Error Details: {error_msg}")
                return None, f"خطأ المحرك (Code {process.returncode}): {error_msg}"
    except Exception as e:
        logger.error(f"❌ Local LibreOffice Exception: {str(e)}")
        return None, f"استثناء المحرك: {str(e)}"

# 💡 الرادار اللغوي الذكي (يحدد هل النص عربي أم لاتيني)
def has_arabic(text):
    return bool(re.search(r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]', text))

def get_style_prompt(style, mode):
    global_rules = """
⚠️ DRAFTING VS. FORMATTING & ZERO HALLUCINATION (CRITICAL RULE):
- DRAFTING MODE: If the user asks you to "write", "compose", or "draft" a document based on a brief topic, act as a professional copywriter to structure the letter/document. HOWEVER, YOU MUST STRICTLY USE ONLY THE INFORMATION PROVIDED BY THE USER. DO NOT invent or hallucinate fake names, fake phone numbers, fake prices, or fake company names.
- 🚫 NO PLACEHOLDERS: If the user does not provide a specific piece of information, DO NOT create empty placeholders UNLESS the user explicitly requests a fillable form or blank template. Otherwise, simply OMIT that element entirely.
- FORMATTING MODE: If the user provides ready-made text, a draft, or specific data, your ONLY job is to format their EXACT text into professional HTML. You MUST NOT add, modify, or remove a single word of their content. ZERO hallucination.

⚠️ EXCLUSION RULE:
- 🚫 You MUST completely IGNORE, DELETE, and EXCLUDE any letterheads (headers at the top), footers (at the bottom), logos, stamps, and signatures from the user's uploaded images. DO NOT CREATE FAKE LETTERHEADS.

⚠️ SMART HTML STRUCTURE, DYNAMIC ALIGNMENT & TABLE USAGE (ALL LANGUAGES):
1. 🚫 BREAK THE WALL OF TEXT! A document where every single line starts from the same edge is ugly and rejected. You MUST vary the alignment using inline CSS to make it look like a real printed professional document in ANY language:
   - MAIN TITLES: MUST be strictly centered using `<h1 style="text-align: center;">`.
   - RECIPIENT BLOCK (المرسل إليه / À Monsieur): 
     * For Latin/French/English (e.g., "À monsieur..."): It MUST be aligned to the EXTREME LEFT (`text-align: left; margin-left: 0;`) OR explicitly CENTERED. NEVER push it to the right.
     * For Arabic (e.g., "إلى السيد..."): It MUST be strictly aligned to the EXTREME RIGHT (`text-align: right; margin-right: 0;`). NEVER push it to the left or center it.
   - LONG PARAGRAPHS: MUST be justified to fill the line evenly using `text-align: justify;`.
   - METADATA & FILLABLE FIELDS: Place them intelligently (e.g., using flexbox). If a field is meant to be filled by hand after printing (e.g., "التاريخ:" or "التوقيع:"), you MUST leave sufficient physical writing space next to it (e.g., using "التاريخ: ....................") and NEVER push it to the absolute edge of the page.
   - SIGNATURES / SENDER INFO: Move them to the opposite bottom corner or center them. Do not stack everything on one side.
   - Make the layout dynamic, breathing, and visually balanced!
2. AVOID UNNECESSARY TABLES: DO NOT use tables for standard paragraphs, headers, dates, or signatures. Use tables ONLY for actual tabular data grids (e.g., items, prices, schedules).
3. NO DIV TABLES: When you DO use tables, use classical HTML `<table>`, `<tr>`, `<td>`, `<th>`.
4. 🚫 NO GHOST BOXES: NEVER use CSS `border`, `outline`, or `background` on `<div>`, `<p>`, or `<span>`. Borders are STRICTLY allowed ONLY on `<table>`, `<th>`, and `<td>`.
5. 🚫 NO EMPTY ROWS: NEVER create empty `<tr>` rows or spacer rows at the top of the table. Start directly with the actual text headers. Do NOT use `<thead>`, `<tbody>`, or `<tfoot>` tags.
6. 📊 INVOICE TOTALS (COLSPAN): For rows calculating "Total" (الإجمالي), use the `colspan` attribute to merge empty cells nicely.
7. 🚫 NO FIXED FONT SIZES (CRITICAL FOR DYNAMIC SCALING): NEVER use hardcoded pixel or point sizes (e.g., `font-size: 14px;` or `12pt`). You MUST use standard semantic tags (`<h1>`, `<h2>`) or relative sizes (e.g., `font-size: 1.2em;`, `120%`) for visual hierarchy. The parent container controls the base size.

⚠️ BIDI & MULTILINGUAL LAYOUT LOCKS (MANDATORY):
- Outermost wrapper MUST use `dir="ltr"`.
- Arabic text and `<table>` elements MUST explicitly use `dir="rtl"`.
- Latin/French/English text and `<table>` elements MUST explicitly use `dir="ltr"`.
- 🔄 TABLE COLUMN ORDER: Extract columns in their exact natural logical order as they appear. DO NOT attempt to manually reverse or flip the columns. The browser handles RTL/LTR table rendering automatically based on the `dir` attribute.
- NUMBER ANTI-REVERSAL: ALL numbers MUST strictly be wrapped in: `<span dir="ltr" style="display:inline-block; direction:ltr; unicode-bidi:isolate; white-space:nowrap;"></span>`.
"""
    if mode == "simulation":
        return f"""CLONING: Reproduce EXACTLY text/tables from the reference image.
IGNORE logos, stamps, signatures. Do NOT invent data.
⚠️ EXCEPTIONAL SCENARIO: If the image is a SINGLE circular stamp, produce ONLY an inline <svg> element.
{global_rules}
RULE E – NO BORDERS: You MUST NOT add any outer border, stroke, or page-like box.
RULE F – CAMERA DISTORTION: Ignore physical distortion. Reconstruct in its NATURAL format adapting to the canvas."""

    design_base = ""
    if style == "modern":
        design_base = """MODERN/CREATIVE - Professional, beautiful, and highly aesthetic document design.
CREATIVE FREEDOM: Choose harmonious modern color palettes, elegant typography. Use soft background colors for table headers."""
    else:
        design_base = """FORMAL/OFFICIAL - Ultra clean, strictly official document design.
⚠️ CRITICAL HEADINGS RULE: ABSOLUTELY NO vertical lines, NO border-left, NO border-right, and NO blockquotes next to any headings. Headings MUST be plain, clean, bold text.
⚠️ CRITICAL TABLE RULE: STRICTLY use plain `<table>` with pure black borders. NO background colors, NO gray cells, NO shaded rows. Keep it 100% formal, printable, and transparent.
TYPOGRAPHY: Dynamic sizes. Title bold centered."""

    return f"{design_base}\n\n{global_rules}"


def detect_document_type(user_msg):
    msg_lower = user_msg.lower()
    single_page_keywords = ['فاتورة', 'facture', 'invoice', 'devis', 'عرض سعر', 'bon', 'شهادة', 'certificate', 'attestation', 'رسالة', 'letter', 'lettre', 'courrier', 'إيصال', 'receipt', 'reçu', 'تصريح', 'declaration', 'إذن', 'autorisation', 'بطاقة', 'card']
    multi_page_keywords = ['تقرير', 'report', 'rapport', 'دراسة', 'study', 'étude', 'بحث', 'research', 'خطة', 'plan', 'مشروع', 'project', 'تفصيلي', 'detailed', 'شامل', 'comprehensive']
    for kw in single_page_keywords:
        if kw in msg_lower: return "single_page"
    for kw in multi_page_keywords:
        if kw in msg_lower: return "multi_page"
    return "auto"


@app.route("/", methods=["GET"])
def index():
    return jsonify({"status": "Monjez V10 Server Active", "features": ["documents", "simulation", "design", "translation", "word_export", "magic_convert"]})

@app.route("/gemini", methods=["POST"])
def generate():
    if not get_client(): return jsonify({"error": "Gemini API Offline"}), 500
    try:
        data = request.json
        user_msg = data.get("message", "")
        mode = data.get("mode", "documents")
        style = data.get("style", "formal")
        page_size = data.get("pageSize", "a4Portrait")
        reference_b64 = data.get("reference_image")
        letterhead_b64 = data.get("letterhead_image")

        style_prompt = get_style_prompt(style, mode)
        doc_type = detect_document_type(user_msg)

        page_dimensions = {
            "a4Portrait": {"w": 595, "h": 842, "orientation": "portrait", "physical": "21.0cm x 29.7cm"},
            "a4Landscape": {"w": 842, "h": 595, "orientation": "landscape", "physical": "29.7cm x 21.0cm"},
            "a3": {"w": 842, "h": 1191, "orientation": "portrait A3", "physical": "29.7cm x 42.0cm"},
            "a5": {"w": 420, "h": 595, "orientation": "portrait A5", "physical": "14.8cm x 21.0cm"},
        }
        page_info = page_dimensions.get(page_size, page_dimensions["a4Portrait"])
        is_landscape = page_info["w"] > page_info["h"]

        landscape_extra = f" LANDSCAPE LAYOUT: Tables MUST fit within this width horizontally, but text can flow naturally downwards." if is_landscape else ""
        orientation_instruction = f"PAGE FORMAT: {page_info['orientation']} — Physical Canvas Size: {page_info['physical']} (Target width: {page_info['w']}px). {landscape_extra}"
        
        ref_note = "\nATTACHED IMAGE: Insert using <img src='data:image/jpeg;base64,...' style='max-width:80%; height:auto; margin:8px auto; display:block;' />" if reference_b64 and mode != "simulation" else ""

        doc_type_instruction = "SINGLE-PAGE DOCUMENT: Optimize space beautifully on one page." if doc_type == "single_page" else "MULTI-PAGE DOCUMENT: Allow natural flow across multiple pages."

        svg_rule = "NO `<html>`, `<body>`. (EXCEPTION: `<svg>` is ONLY allowed for the standalone circular stamp scenario)." if mode == "simulation" else "NO `<svg>`, `<html>`, `<body>`."

        prompt = f"""You are a STRICT Document Formatter.
{style_prompt}
{orientation_instruction}
{ref_note}
{doc_type_instruction}
TECHNICAL RULES:
1. PURE HTML ONLY. Just `<div>`, `<table>`, `<h1>`, `<p>`. {svg_rule}
2. NO BORDERS AROUND DOCUMENT.
3. WRAPPER CONFIG: The outermost wrapper MUST NOT have excessive padding. Use `<div style="width:100%; max-width:100%; margin:0 auto; padding:5px; box-sizing:border-box; direction:ltr; overflow-wrap:anywhere; word-break:break-word; overflow:hidden;">`.
OUTPUT: Return raw HTML only."""

        contents = [user_msg] if user_msg else ["Create a formal document."]
        if reference_b64:
            contents.append(get_types().Part.from_bytes(data=base64.b64decode(reference_b64), mime_type="image/jpeg"))
        if letterhead_b64:
            contents.append("Ensure layout fits empty space below this letterhead.")
            contents.append(get_types().Part.from_bytes(data=base64.b64decode(letterhead_b64), mime_type="image/jpeg"))

        gen_config = get_types().GenerateContentConfig(system_instruction=prompt, temperature=0.15, max_output_tokens=20000)

        try:
            resp = call_gemini("gemini-3-flash-preview", contents, gen_config, 55)
        except:
            resp = call_gemini("gemini-2.5-flash", contents, gen_config, 50)

        clean_html = clean_html_output(resp.text or "")
        used_tokens = extract_tokens(resp)
        logger.info(f"✅ Generated HTML (mode: {mode}, page: {page_size}) | Tokens: {used_tokens}")
        return jsonify({"response": clean_html, "used_tokens": used_tokens})
    except Exception as e:
        logger.error(f"Error: {str(e)}", exc_info=True)
        return jsonify({"error": "Failed", "details": str(e), "used_tokens": 0}), 500


@app.route("/modify", methods=["POST"])
def modify():
    if not get_client(): return jsonify({"error": "Gemini API Offline"}), 500
    try:
        data = request.json
        current_html = data.get("current_html") or data.get("currentSVG") or data.get("current_svg") or data.get("htmlContent") or ""
        instruction = data.get("instruction") or data.get("prompt") or ""
        ref_b64 = data.get("reference_image")

        if not current_html.strip():
            logger.error("❌ ERROR: current_html is empty!")
            return jsonify({"error": "Failed", "details": "لم يتم العثور على محتوى المستند الحالي لإجراء التعديل الذكي. يرجى المحاولة مرة أخرى."}), 400

        img_note = f"\nINSERT image: <img src='data:image/jpeg;base64,{ref_b64}' style='max-width:80%; height:auto; margin:8px auto; display:block;' />" if ref_b64 else ""

        sys = f"""You are a STRICT HTML PATCHING ENGINE. You are NOT a designer.
You will receive a <CURRENT_HTML> document and a <USER_REQUEST>.

CRITICAL RULES (MUST FOLLOW STRICTLY):
1. EXACT COPY-PASTE: Output the EXACT SAME HTML structure provided. DO NOT delete unrelated text or sections.
2. SURGICAL EDIT: Apply the exact surgical change requested. DO NOT hallucinate or add fake elements.
3. BIDI & TYPOGRAPHY PROTECTION: 
   - Preserve `dir="ltr"` on wrappers. Arabic `<table>` elements use `dir="rtl"`.
   - Protect phone numbers with `<span dir="ltr" style="display:inline-block; unicode-bidi:bidi-override; white-space:nowrap;">`.
   - Text in Arabic MUST use `font-family: 'Arial', sans-serif;`. Text in Latin/English MUST use `font-family: 'Times New Roman', serif;`.
4. 🚫 NO BORDERS & NO BACKGROUNDS (CRITICAL): NEVER add outer borders, strokes, shadow boxes, or background colors (especially dark ones) to the main wrappers (`<div>`, `<p>`, `<span>`). The document MUST remain a clean, borderless, transparent standard paper layout.
5. NO FORCED SPACING: DO NOT inject inline `line-height` or custom `margin/padding` into text elements (`<p>`, `<span>`) unless the user explicitly asks for it. Rely on the document's global layout.
6. RETURN FULL HTML: Return the complete patched HTML. Do not truncate or use placeholders like ''.
{img_note}

OUTPUT FORMAT:
[MESSAGE]
وصف قصير للتعديل باللغة العربية
[/MESSAGE]
[HTML]
(ضع هنا كود الـ HTML المعدل كاملاً)
[/HTML]"""

        cfg = get_types().GenerateContentConfig(
            system_instruction=sys, 
            temperature=0.0, 
            max_output_tokens=16384
        )

        cts = [f"<CURRENT_HTML>\n{current_html}\n</CURRENT_HTML>\n\n<USER_REQUEST>\n{instruction}\n</USER_REQUEST>\n\nTASK: Apply the exact surgical change and return FULL updated HTML."]
        if ref_b64:
            cts.append(get_types().Part.from_bytes(data=base64.b64decode(ref_b64), mime_type="image/jpeg"))

        try:
            resp = call_gemini("gemini-3-flash-preview", cts, cfg, 55)
        except:
            resp = call_gemini("gemini-2.5-flash", cts, cfg, 50)

        used_tokens = extract_tokens(resp)
        text = resp.text or ""
        msg_match = re.search(r'\[MESSAGE\](.*?)\[/MESSAGE\]', text, re.DOTALL | re.IGNORECASE)
        html_match = re.search(r'\[HTML\](.*?)\[/HTML\]', text, re.DOTALL | re.IGNORECASE)

        if html_match:
            new_inner = html_match.group(1).strip()
            message = msg_match.group(1).strip() if msg_match else "تم التعديل بنجاح ✨"
        else:
            new_inner = clean_html_output(text)
            new_inner = re.sub(r'\[MESSAGE\].*?\[/MESSAGE\]', '', new_inner, flags=re.DOTALL | re.IGNORECASE).strip()
            message = "تم التعديل بنجاح ✨"

        return jsonify({"response": new_inner, "message": message, "used_tokens": used_tokens})
    except Exception as e:
        logger.error(f"Modify Error: {str(e)}", exc_info=True)
        return jsonify({"error": "Failed", "details": str(e), "used_tokens": 0}), 500


@app.route("/format", methods=["POST"])
def smart_format():
    if not get_client(): return jsonify({"error": "Gemini API Offline"}), 500
    try:
        data = request.json
        current_html = data.get("current_html", "")

        sys = f"""You are a STRICT Document Editor. The user has manually edited this document.
YOUR MISSION:
1. CLEANUP & STRUCTURE: Wrap loose text in proper tags. Apply logical Alignments.
2. STRICT PRESERVATION: NEVER delete, alter, or add to the actual facts, text, or meaning. NO HALLUCINATION.
3. BIDI FIX: Ensure wrappers use `dir="ltr"`. Arabic `<table>` elements use `dir="rtl"`. Arabic text uses `dir="rtl" style="text-align:right"`. Protect phone numbers.
OUTPUT FORMAT:
[MESSAGE]
تم تنسيق وترتيب المستند بنجاح ✨
[/MESSAGE]
[HTML]
(ضع هنا كود الـ HTML المنسق كاملاً)
[/HTML]"""

        cfg = get_types().GenerateContentConfig(system_instruction=sys, temperature=0.0, max_output_tokens=16384)
        cts = [f"<MESSY_HTML>\n{current_html}\n</MESSY_HTML>\n\nPlease format and fix Bidi issues professionally without changing text."]

        try:
            resp = call_gemini("gemini-3-flash-preview", cts, cfg, 55)
        except:
            resp = call_gemini("gemini-2.5-flash", cts, cfg, 50)

        used_tokens = extract_tokens(resp)
        text = resp.text or ""
        msg_match = re.search(r'\[MESSAGE\](.*?)\[/MESSAGE\]', text, re.DOTALL | re.IGNORECASE)
        html_match = re.search(r'\[HTML\](.*?)\[/HTML\]', text, re.DOTALL | re.IGNORECASE)

        if html_match:
            new_inner = html_match.group(1).strip()
            message = msg_match.group(1).strip() if msg_match else "تم التنسيق ✨"
        else:
            new_inner = clean_html_output(text)
            new_inner = re.sub(r'\[MESSAGE\].*?\[/MESSAGE\]', '', new_inner, flags=re.DOTALL | re.IGNORECASE).strip()
            message = "تم التنسيق ✨"

        logger.info("✅ Document Smartly Formatted")
        return jsonify({"response": new_inner, "message": message, "used_tokens": used_tokens})
    except Exception as e:
        logger.error(f"Format Error: {str(e)}", exc_info=True)
        return jsonify({"error": "Failed", "details": str(e), "used_tokens": 0}), 500


# ══════════════════════════════════════════════════════════
# مسار تحويل HTML/PDF إلى Word 
# ══════════════════════════════════════════════════════════
@app.route("/convert_to_word", methods=["POST"])
def convert_to_word():
    try:
        data = request.json
        html_content = data.get("html_content") or data.get("current_html")
        pdf_b64 = data.get("pdf_base64", "")
        letterhead_b64 = data.get("letterhead_base64", "") 
        letterhead_on_all_pages = data.get("letterhead_on_all_pages", False)
        
        input_fmt = "html"
        used_tokens = 0

        if pdf_b64 and not html_content:
            logger.info("📄 Converting PDF to Word via AI Bridge first (To preserve tables)...")
            gemini_bytes = base64.b64decode(pdf_b64)
            
            bridge_prompt = """You are an OCR and Document Extraction Engine.
Your task is to precisely extract ALL content from the attached document and convert it into a fully structured, professional HTML document.
CRITICAL RULES:
1. NO HALLUCINATIONS: Extract the exact words, numbers, and tables. Do not summarize or invent text.
2. 🚫 CRITICAL EXCLUSION RULE: IGNORE, DELETE, and EXCLUDE any letterheads, footers, logos, stamps, and signatures.
3. TABLES & COLSPAN: Use proper `<table>`. NO background colors. For "Total" (الإجمالي) rows, use `colspan` nicely.
4. 🚫 NO EMPTY ROWS: NEVER create empty `<tr>` rows or spacer cells. Do NOT use `<thead>`, `<tbody>`, or `<tfoot>` tags.
5. 🚫 NO GHOST BOXES: NEVER use CSS borders on `<div>`, `<p>`, or `<span>`.
6. 🔄 COLUMN ORDER: Extract columns exactly as they appear in their natural logical order without reversing them.
7. NUMBERS: Wrap any standalone numbers/dates in `<span dir="ltr"></span>`.
8. NO MARKDOWN: Output strictly pure HTML code."""
            
            contents = [bridge_prompt, get_types().Part.from_bytes(data=gemini_bytes, mime_type="application/pdf")]
            gen_config = get_types().GenerateContentConfig(temperature=0.0, max_output_tokens=16384)
            
            try: resp = call_gemini("gemini-3-flash-preview", contents, gen_config, 90)
            except: resp = call_gemini("gemini-2.5-flash", contents, gen_config, 90)
            
            used_tokens = extract_tokens(resp)
            extracted_html = clean_html_output(resp.text or "")
            if not extracted_html:
                return jsonify({"error": "Failed", "details": "فشل الذكاء الاصطناعي في قراءة الـ PDF.", "used_tokens": used_tokens}), 500
            
            html_content = extracted_html

        if html_content:
            logger.info("📄 Preparing HTML for LibreOffice Word Conversion...")

            html_content = force_table_borders(html_content)
            html_content = force_tables_ltr_for_export(html_content)
            html_content = re.sub(r'font-family\s*:[^;"]+[;]?', '', html_content, flags=re.IGNORECASE)
            
            # 💡 لحام الأرقام لمنع انعكاسها
            html_content = re.sub(r'(\d)\s+(?=\d)', r'\1&nbsp;', html_content)
            
            is_arabic_doc = has_arabic(html_content)
            body_dir = "rtl" if is_arabic_doc else "ltr"

            html_content = re.sub(
                r'<div[^>]*display\s*:\s*flex[^>]*>.*?<div[^>]*border-bottom[^>]*>.*?</div>.*?<div[^>]*>\s*:\s*</div>.*?<div[^>]*>(.*?)</div>.*?</div>',
                r'<p dir="rtl" style="text-align:right; margin:0;">\1: ........................................</p>',
                html_content, flags=re.IGNORECASE | re.DOTALL)
            html_content = re.sub(
                r'<div[^>]*display\s*:\s*flex[^>]*>.*?<div[^>]*>(.*?)</div>.*?<div[^>]*>\s*:\s*</div>.*?<div[^>]*border-bottom[^>]*>.*?</div>.*?</div>',
                r'<p dir="rtl" style="text-align:right; margin:0;">\1: ........................................</p>',
                html_content, flags=re.IGNORECASE | re.DOTALL)
            html_content = re.sub(r'<div[^>]*border-bottom[^>]*>(\s|&nbsp;)*</div>', ' ........................................ ', html_content, flags=re.IGNORECASE)

            full_html = f"""<html lang="ar" dir="{body_dir}">
<head>
<meta charset="utf-8">
<style>
  * {{ font-family: 'Arial', sans-serif !important; }}
  table {{ border-collapse: collapse; margin: 10px 0; width: 100% !important; }}
  th, td {{ border: 1px solid #000000; padding: 4px !important; line-height: 1.1 !important; vertical-align: middle !important; }}
  p, h1, h2, h3, h4, h5, h6, div, span {{ margin: 0; padding: 0; border: none !important; background: transparent !important; line-height: 1.2 !important; }}
</style>
</head>
<body>
{html_content}
</body>
</html>"""
            file_bytes = full_html.encode('utf-8')
            
        else:
            return jsonify({"error": "Failed", "details": "لم يتم إرسال محتوى المستند.", "used_tokens": 0}), 400

        raw_docx_bytes, err_msg = local_libreoffice_convert(file_bytes, input_fmt, "docx")
        
        if not raw_docx_bytes:
            return jsonify({"error": "Failed", "details": f"فشل LibreOffice: {err_msg}", "used_tokens": used_tokens}), 500

        # ══════════════════════════════════════════════════════════
        # معالجة الوورد: الحرية للمحاذاة وإصلاح انعكاس الأعمدة والتاريخ
        # ══════════════════════════════════════════════════════════
        logger.info("💉 Local Processing: Deep XML Fixes for Fonts and Tables...")
        doc_stream = io.BytesIO(raw_docx_bytes)
        doc = docx.Document(doc_stream)
        section = doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Cm(4.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

        def clean_and_format_paragraph(paragraph, is_table=False):
            text = paragraph.text.strip()
            is_arabic_para = has_arabic(text) if text else is_arabic_doc

            pPr = paragraph._element.get_or_add_pPr()
            
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = OxmlElement('w:spacing')
                pPr.append(spacing)
            
            # ✅ ضبط التباعد (مرونة للنص العادي، وتماسك للجداول)
            if is_table:
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '60')  # تباعد خفيف للجدول
                spacing.set(qn('w:line'), '280')  
            else:
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '240') # تباعد واضح بين الفقرات (12pt)
                spacing.set(qn('w:line'), '360')  # تباعد أسطر يعادل 1.5 للنصوص العادية
            spacing.set(qn('w:lineRule'), 'auto')

            # ✅ ضبط حجم الخط (30 تعني 15pt للنص العادي | 24 تعني 12pt للجدول)
            target_size = '24' if is_table else '30'

            for run in paragraph.runs:
                rPr = run._element.get_or_add_rPr()
                
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:cs'), 'Arial')
                rFonts.set(qn('w:ascii'), 'Arial')
                rFonts.set(qn('w:hAnsi'), 'Arial')
                rFonts.set(qn('w:eastAsia'), 'Arial')
                
                sz = rPr.find(qn('w:sz'))
                if sz is None: sz = OxmlElement('w:sz'); rPr.append(sz)
                sz.set(qn('w:val'), target_size)
                
                szCs = rPr.find(qn('w:szCs'))
                if szCs is None: szCs = OxmlElement('w:szCs'); rPr.append(szCs)
                szCs.set(qn('w:val'), target_size)

        for table in doc.tables:
            table.autofit = True
            tblPr = table._element.tblPr
            if tblPr is not None:
                tblLayout = tblPr.find(qn('w:tblLayout'))
                if tblLayout is not None:
                    tblLayout.set(qn('w:type'), 'autofit')
                
                tblW = tblPr.find(qn('w:tblW'))
                if tblW is None:
                    tblW = OxmlElement('w:tblW')
                    tblPr.append(tblW)
                tblW.set(qn('w:w'), '5000') 
                tblW.set(qn('w:type'), 'pct')

                # 💡 [الحل الجذري المعتمد]: إجبار الجدول على الاتجاه العربي في الوورد ليمنع انعكاس الأعمدة
                if is_arabic_doc:
                    bidiVisual = tblPr.find(qn('w:bidiVisual'))
                    if bidiVisual is None:
                        bidiVisual = OxmlElement('w:bidiVisual')
                        tblPr.append(bidiVisual)

            for row in table.rows:
                trPr = row._tr.get_or_add_trPr()
                for trHeight in trPr.findall(qn('w:trHeight')):
                    trPr.remove(trHeight)
                
                for cell in row.cells:
                    tcPr = cell._element.get_or_add_tcPr()
                    
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is not None:
                        tcW.set(qn('w:type'), 'auto')

                    tcMar = tcPr.find(qn('w:tcMar'))
                    if tcMar is not None: tcPr.remove(tcMar)
                    tcMar = OxmlElement('w:tcMar')
                    for attr in ['top', 'bottom']:
                        node = OxmlElement(f'w:{attr}')
                        node.set(qn('w:w'), '0')
                        node.set(qn('w:type'), 'dxa')
                        tcMar.append(node)
                    tcPr.append(tcMar)

                    for p in cell.paragraphs:
                        clean_and_format_paragraph(p, is_table=True)

        for p in doc.paragraphs:
            clean_and_format_paragraph(p, is_table=False)

        if letterhead_b64:
            header_img_data = base64.b64decode(letterhead_b64)
            header_img_stream = io.BytesIO(header_img_data)
            
            if not letterhead_on_all_pages:
                section.different_first_page_header_footer = True
                target_header = section.first_page_header
            else:
                target_header = section.header
                
            if not target_header.paragraphs:
                target_header.add_paragraph()
            paragraph = target_header.paragraphs[0]
            
            run = paragraph.add_run()
            shape = run.add_picture(header_img_stream, width=Inches(8.27), height=Inches(11.69))
            
            inline = shape._inline
            anchor = OxmlElement('wp:anchor')
            anchor.set('distT', '0'); anchor.set('distB', '0'); anchor.set('distL', '0'); anchor.set('distR', '0')
            anchor.set('simplePos', '0'); anchor.set('relativeHeight', '0'); anchor.set('behindDoc', '1') 
            anchor.set('locked', '0'); anchor.set('layoutInCell', '1'); anchor.set('allowOverlap', '1')

            simplePos = OxmlElement('wp:simplePos')
            simplePos.set('x', '0'); simplePos.set('y', '0')
            anchor.append(simplePos)

            positionH = OxmlElement('wp:positionH')
            positionH.set('relativeFrom', 'page')
            alignH = OxmlElement('wp:align'); alignH.text = 'center'
            positionH.append(alignH); anchor.append(positionH)
            
            positionV = OxmlElement('wp:positionV')
            positionV.set('relativeFrom', 'page')
            alignV = OxmlElement('wp:align'); alignV.text = 'top'
            positionV.append(alignV); anchor.append(positionV)
            
            anchor.append(inline.extent)
            effectExtent = OxmlElement('wp:effectExtent')
            effectExtent.set('l', '0'); effectExtent.set('t', '0'); effectExtent.set('r', '0'); effectExtent.set('b', '0')
            anchor.append(effectExtent)
            
            anchor.append(OxmlElement('wp:wrapNone'))
            anchor.append(inline.docPr)
            anchor.append(OxmlElement('wp:cNvGraphicFramePr'))
            anchor.append(inline.graphic)
            
            inline.getparent().replace(inline, anchor)

        final_docx_stream = io.BytesIO()
        doc.save(final_docx_stream)
        docx_bytes = final_docx_stream.getvalue()
        docx_b64 = base64.b64encode(docx_bytes).decode('utf-8')

        logger.info(f"✅ Final Word Document generated successfully ({len(docx_bytes)} bytes)")
        return jsonify({"docx_base64": docx_b64, "message": "تم التحويل إلى Word بنجاح ✨", "used_tokens": used_tokens})
    except Exception as e:
        logger.error(f"Word Error: {str(e)}", exc_info=True)
        return jsonify({"error": "Failed", "details": f"فشل التحويل: {str(e)}", "used_tokens": 0}), 500


# ══════════════════════════════════════════════════════════
# مسار MAGIC CONVERTER (المحول الشامل)
# ══════════════════════════════════════════════════════════
@app.route("/magic_convert", methods=["POST"])
def magic_convert():
    try:
        data = request.json
        file_b64 = data.get("fileBase64")
        mime_type = data.get("mimeType", "")
        target_format = data.get("targetFormat", "word")
        is_arabic = data.get("isArabic", True)
        extract_only = data.get("extractOnly", False)  

        if not file_b64:
            return jsonify({"error": "Failed", "details": "لم يتم العثور على الملف", "used_tokens": 0}), 400

        mime_lower = mime_type.lower()
        input_ext = "pdf"
        if "html" in mime_lower: input_ext = "html"
        elif "jpeg" in mime_lower or "jpg" in mime_lower: input_ext = "jpg"
        elif "png" in mime_lower: input_ext = "png"
        elif "msword" in mime_lower or "word" in mime_lower or "docx" in mime_lower: input_ext = "docx" 
        elif "excel" in mime_lower or "xls" in mime_lower or "spreadsheet" in mime_lower: input_ext = "xlsx"
        elif "powerpoint" in mime_lower or "ppt" in mime_lower or "presentation" in mime_lower: input_ext = "pptx"
        
        file_bytes = base64.b64decode(file_b64)
        
        output_ext = "docx"
        if target_format == "excel": output_ext = "xlsx"
        elif target_format == "powerpoint": output_ext = "pptx"
        elif target_format == "pdf": output_ext = "pdf"
        elif target_format == "html": output_ext = "html"

        logger.info(f"🔄 Magic Request: {input_ext.upper()} ➡️ {output_ext.upper()}")

        direct_conversions = [
            ("docx", "pdf"), ("doc", "pdf"),
            ("xlsx", "pdf"), ("xls", "pdf"),
            ("pptx", "pdf"), ("ppt", "pdf"),
            ("html", "docx"), ("html", "xlsx"), ("html", "pdf")
        ]
        
        used_tokens = 0

        if (input_ext, output_ext) in direct_conversions and not extract_only:
            logger.info("⚡ Route 1: Direct LibreOffice Conversion (No AI needed)...")
            
            if input_ext == "html":
                html_text = file_bytes.decode('utf-8')
                html_text = force_table_borders(html_text)
                html_text = force_tables_ltr_for_export(html_text)
                html_text = re.sub(r'font-family\s*:[^;"]+[;]?', '', html_text, flags=re.IGNORECASE)
                
                html_text = re.sub(r'(\d)\s+(?=\d)', r'\1&nbsp;', html_text)
                is_arabic_doc = has_arabic(html_text)
                body_dir = "rtl" if is_arabic_doc else "ltr"
                
                full_html = f"""<html lang="ar" dir="{body_dir}"><head><meta charset="utf-8">
<style>* {{ font-family: 'Arial', sans-serif !important; }} table {{ border-collapse: collapse; margin: 10px 0; width: 100% !important; }} th, td {{ border: 1px solid #000; padding: 4px !important; line-height: 1.1 !important; }} p, h1, h2, h3, h4, h5, h6, div, span {{ margin: 0; padding: 0; border: none !important; background: transparent !important; }}</style>
</head><body>{html_text}</body></html>"""
                file_bytes = full_html.encode('utf-8')

            result_bytes, err_msg = local_libreoffice_convert(file_bytes, input_ext, output_ext)
            
            if result_bytes:
                result_b64 = base64.b64encode(result_bytes).decode('utf-8')
                return jsonify({
                    "file_base64": result_b64,
                    "extension": output_ext,
                    "message": f"تم التحويل إلى {target_format.upper()} بنجاح ✨",
                    "used_tokens": used_tokens
                })
            else:
                logger.warning(f"⚠️ Direct conversion failed: {err_msg}. Falling back to AI Route if applicable.")

        logger.info("🧠 Route 2: AI OCR & Extraction Bridge...")
        gemini_bytes = file_bytes
        gemini_mime = "application/pdf"
        
        if input_ext in ["docx", "doc", "xlsx", "xls", "pptx", "ppt"]:
            logger.info("🔄 Converting Document to PDF first via LibreOffice for AI Reading...")
            gemini_bytes, err_msg = local_libreoffice_convert(file_bytes, input_ext, "pdf")
            if not gemini_bytes:
                return jsonify({"error": "Failed", "details": f"فشل تجهيز المستند للقراءة: {err_msg}", "used_tokens": 0}), 500
            gemini_mime = "application/pdf"
        elif input_ext in ["jpg", "jpeg"]: gemini_mime = "image/jpeg"
        elif input_ext == "png": gemini_mime = "image/png"

        target_focus = "tables and grids format specifically for Excel" if output_ext == "xlsx" else "general document structure"
        
        bridge_prompt = f"""You are an elite OCR and Document Extraction Engine.
Your task is to precisely extract ALL content from the attached document and convert it into a fully structured, professional HTML document. Focus on {target_focus}.

CRITICAL RULES:
1. NO HALLUCINATIONS: Extract the exact words, numbers, and tables. Do not summarize or invent text.
2. 🚫 CRITICAL EXCLUSION RULE: IGNORE, DELETE, and EXCLUDE any letterheads, footers, logos, stamps, and signatures.
3. TABLES FOR GRIDS ONLY: Use `<table>` ONLY for actual tabular data (items, prices, schedules). Regular text, headers, and dates MUST be in `<p>` or `<div>`. NEVER put the whole document in a table.
4. COLSPAN: For "Total" (الإجمالي) rows, use `colspan` elegantly.
5. 🚫 NO EMPTY ROWS: NEVER create empty `<tr>` or `<th>` rows. Start directly with the text headers. Do NOT use `<thead>`, `<tbody>`, or `<tfoot>` tags.
6. 🚫 NO GHOST BOXES: NEVER use CSS borders on `<div>`, `<p>`, or `<span>`. Borders are for tables ONLY.
7. 🔄 COLUMN ORDER: Extract columns exactly as they appear in their natural logical order without reversing them.
8. NUMBERS: Wrap standalone numbers/dates in `<span dir="ltr"></span>`.
9. PURE HTML ONLY. Do not wrap in ```html."""
        
        contents = [bridge_prompt, get_types().Part.from_bytes(data=gemini_bytes, mime_type=gemini_mime)]
        gen_config = get_types().GenerateContentConfig(temperature=0.0, max_output_tokens=16384)
        
        try: resp = call_gemini("gemini-3-flash-preview", contents, gen_config, 90)
        except: resp = call_gemini("gemini-2.5-flash", contents, gen_config, 90)
        
        used_tokens = extract_tokens(resp)
        extracted_html = clean_html_output(resp.text or "")
        if not extracted_html:
            return jsonify({"error": "Failed", "details": "فشل الذكاء الاصطناعي في قراءة الملف", "used_tokens": used_tokens}), 500
        
        if extract_only or target_format == "html":
            return jsonify({"html_content": extracted_html, "message": "تم استخراج النصوص بنجاح ✨", "used_tokens": used_tokens})
        
        logger.info(f"📄 Wrapping extracted HTML to final format: {output_ext.upper()}...")
        
        extracted_html = force_table_borders(extracted_html)
        extracted_html = force_tables_ltr_for_export(extracted_html)
        extracted_html = re.sub(r'(\d)\s+(?=\d)', r'\1&nbsp;', extracted_html)
        
        is_arabic_doc = has_arabic(extracted_html)
        body_dir = "rtl" if is_arabic_doc else "ltr"
        
        full_html = f"""<html lang="ar" dir="{body_dir}"><head><meta charset="utf-8">
<style>* {{ font-family: 'Arial', sans-serif !important; }} table {{ border-collapse: collapse; margin: 10px 0; width: 100% !important; }} th, td {{ border: 1px solid #000; padding: 4px !important; line-height: 1.1 !important; }} p, h1, h2, h3, h4, h5, h6, div, span {{ margin: 0; padding: 0; border: none !important; background: transparent !important; line-height: 1.2 !important; }}</style>
</head><body>{extracted_html}</body></html>"""
        
        final_bytes = full_html.encode('utf-8')
        result_bytes, err_msg = local_libreoffice_convert(final_bytes, "html", output_ext)
        
        if not result_bytes:
            return jsonify({"error": "Failed", "details": f"فشل تجميع الملف النهائي: {err_msg}", "used_tokens": used_tokens}), 500
            
        result_b64 = base64.b64encode(result_bytes).decode('utf-8')
        return jsonify({
            "file_base64": result_b64,
            "extension": output_ext,
            "message": f"تم التحويل إلى {target_format.upper()} بنجاح ✨",
            "used_tokens": used_tokens
        })

    except Exception as e:
        logger.error(f"Magic Convert Error: {str(e)}", exc_info=True)
        return jsonify({"error": "Failed", "details": str(e), "used_tokens": 0}), 500


# ══════════════════════════════════════════════════════════
# مسار TRANSLATE DOCUMENT
# ══════════════════════════════════════════════════════════
@app.route("/translate_document", methods=["POST"])
def translate_document():
    if not get_client(): return jsonify({"error": "Gemini API Offline"}), 500
    try:
        data = request.json
        target_language = data.get("target_language", "العربية")
        reference_b64 = data.get("reference_image")
        page_size = data.get("pageSize", "a4Portrait")

        page_dimensions = {
            "a4Portrait": {"w": 595, "h": 842, "orientation": "portrait", "physical": "21.0cm x 29.7cm"},
            "a4Landscape": {"w": 842, "h": 595, "orientation": "landscape", "physical": "29.7cm x 21.0cm"},
            "a3": {"w": 842, "h": 1191, "orientation": "portrait A3", "physical": "29.7cm x 42.0cm"},
            "a5": {"w": 420, "h": 595, "orientation": "portrait A5", "physical": "14.8cm x 21.0cm"},
        }
        page_info = page_dimensions.get(page_size, page_dimensions["a4Portrait"])
        is_landscape = page_info["w"] > page_info["h"]

        landscape_extra = f" LANDSCAPE LAYOUT: Tables MUST fit within this width horizontally, but text can flow naturally downwards." if is_landscape else ""
        orientation_instruction = f"PAGE FORMAT: {page_info['orientation']} — Physical Canvas Size: {page_info['physical']} (Target width: {page_info['w']}px). {landscape_extra}"

        bidi_rules = """
⚠️ BIDI & LAYOUT LOCKS:
- Outermost wrapper MUST use `dir="ltr"`.
- Arabic `<table>` elements MUST use `dir="rtl"`.
- Non-Arabic (Latin/French) `<table>` elements MUST use `dir="ltr"`.
- Arabic text MUST explicitly use `dir="rtl" style="text-align: right;"`
- TABLE COLUMN ORDER: Output HTML columns in their NATURAL logical order exactly as they appear. DO NOT manually reverse the columns.
- NUMBER ANTI-REVERSAL: ALL numbers MUST strictly be wrapped in: `<span dir="ltr" style="display:inline-block; direction:ltr; unicode-bidi:isolate; white-space:nowrap;"></span>`.
"""

        prompt = f"""You are an Expert Professional Translator and Strict Document Formatter.
YOUR MISSION:
1. Clone the exact layout, structure, and tables of the provided document image.
2. TRANSLATE all text into {target_language} with high professional accuracy. 
3. DO NOT invent fake data, logos, or headers. Translate exactly what is there.
4. 🚫 CRITICAL EXCLUSION RULE: You MUST completely IGNORE, DELETE, and EXCLUDE any original letterheads, footers, logos, stamps, and signatures.
{bidi_rules}
{orientation_instruction}
TECHNICAL RULES:
1. PURE HTML ONLY. Just `<div>`, `<table>`, `<h1>`, `<p>`. NO `<svg>`, `<html>`, `<body>`.
2. NO BORDERS AROUND DOCUMENT.
OUTPUT: Return raw HTML only."""

        contents = [f"Translate this document to {target_language} while keeping the exact layout."]
        if reference_b64:
            contents.append(get_types().Part.from_bytes(data=base64.b64decode(reference_b64), mime_type="image/jpeg"))
        else:
            return jsonify({"error": "Failed", "details": "لم يتم إرفاق المستند", "used_tokens": 0}), 400

        gen_config = get_types().GenerateContentConfig(system_instruction=prompt, temperature=0.15, max_output_tokens=20000)

        try:
            resp = call_gemini("gemini-3-flash-preview", contents, gen_config, 55)
        except:
            resp = call_gemini("gemini-2.5-flash", contents, gen_config, 50)

        used_tokens = extract_tokens(resp)
        clean_html = clean_html_output(resp.text or "")
        logger.info(f"✅ Generated Translation HTML (Target: {target_language}) | Tokens: {used_tokens}")
        return jsonify({"response": clean_html, "used_tokens": used_tokens})
    except Exception as e:
        logger.error(f"Error: {str(e)}", exc_info=True)
        return jsonify({"error": "Failed", "details": str(e), "used_tokens": 0}), 500

# ══════════════════════════════════════════════════════════
# 🚀 NEW: DESIGN GENERATION (Vertex AI: Imagen 4 Ultra -> 3 Fallback)
# ══════════════════════════════════════════════════════════
import os
import re
import json
import logging
import base64
import time
import io
import concurrent.futures
import subprocess
import tempfile
from flask import Flask, request, jsonify

# ══════════════════════════════════════════════════════════
# ✅ استدعاء مكتبات الوورد المطلوبة للحقن العميق للرأسية وضبط الهوامش
# ══════════════════════════════════════════════════════════
import docx
from docx.shared import Inches, Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger("Monjez_V10_Server")

app = Flask(__name__)

# ── Lazy Gemini ──
_client = None
_types = None
_init = False

def get_client():
    global _client, _types, _init
    if not _init:
        _init = True
        try:
            from google import genai as g
            from google.genai import types as t
            _types = t
            k = os.environ.get("GOOGLE_API_KEY")
            if k:
                _client = g.Client(api_key=k, http_options={"api_version": "v1beta"})
                logger.info("✅ Monjez V10 Server (Ready)")
        except Exception as e:
            logger.error(f"Init: {e}")
    return _client

def get_types():
    get_client()
    return _types

def call_gemini(model, contents, config, timeout):
    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as ex:
        f = ex.submit(get_client().models.generate_content, model=model, contents=contents, config=config)
        return f.result(timeout=timeout)

# 💡 دالة جديدة لاستخراج الاستهلاك الدقيق للتوكنز
def extract_tokens(resp):
    try:
        if hasattr(resp, 'usage_metadata') and resp.usage_metadata:
            return getattr(resp.usage_metadata, 'total_token_count', 0)
    except Exception as e:
        logger.error(f"Token extraction error: {e}")
    return 0

def clean_html_output(raw_text):
    raw = raw_text.strip()
    if raw.startswith("`" * 3):
        raw = re.sub(r"^`{3}(?:html|xml)?\n?", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\n?`{3}$", "", raw)
    div_match = re.search(r'<div[^>]*xmlns="http://www.w3.org/1999/xhtml"[^>]*>(.*?)</div>\s*</foreignObject>', raw, re.DOTALL)
    if div_match:
        raw = div_match.group(1)
    raw = re.sub(r'\s?contenteditable="[^"]*"', '', raw, flags=re.IGNORECASE)
    raw = re.sub(r'\s?contenteditable=\'[^\']*\'', '', raw, flags=re.IGNORECASE)
    raw = re.sub(r'\s?contenteditable', '', raw, flags=re.IGNORECASE)
    return raw.strip()

# ══════════════════════════════════════════════════════════
# 🛡️ حقنة الجداول (درع الخطوط المزدوجة والصفوف الوهمية)
# ══════════════════════════════════════════════════════════
def force_table_borders(html_text):
    # 0. إزالة أوسمة البنية التي يُنشئها Gemini أحياناً وتسبب صفاً وهمياً في LibreOffice
    html_text = re.sub(r'</?thead[^>]*>', '', html_text, flags=re.IGNORECASE)
    html_text = re.sub(r'</?tbody[^>]*>', '', html_text, flags=re.IGNORECASE)
    html_text = re.sub(r'</?tfoot[^>]*>', '', html_text, flags=re.IGNORECASE)
    html_text = re.sub(r'<colgroup[^>]*>.*?</colgroup>', '', html_text, flags=re.IGNORECASE | re.DOTALL)
    html_text = re.sub(r'<caption[^>]*>.*?</caption>', '', html_text, flags=re.IGNORECASE | re.DOTALL)
    
    # 1. إجبار الجدول على التنسيق النظيف المندمج لمنع الخطوط المزدوجة
    html_text = html_text.replace("<table", "<table border='1' cellpadding='4' cellspacing='0' style='border-collapse:collapse; border-spacing:0; width:100%; border: 1px solid black; margin: 10px 0;' ")
    html_text = html_text.replace("<th", "<th style='border: 1px solid black; padding: 4px; text-align: center; vertical-align: middle; color: black;' ")
    html_text = html_text.replace("<td", "<td style='border: 1px solid black; padding: 4px; vertical-align: middle;' ")
    
    # 2. درع التنظيف: مسح أي صفوف فارغة (Empty Rows) أنشأها الذكاء الاصطناعي وتسبب الخانة الفارغة
    html_text = re.sub(r'<tr>\s*(?:<t[hd][^>]*>\s*</t[hd]>\s*)+</tr>', '', html_text, flags=re.IGNORECASE)
    html_text = re.sub(r'<tr>\s*(?:<t[hd][^>]*>\s*&nbsp;\s*</t[hd]>\s*)+</tr>', '', html_text, flags=re.IGNORECASE)
    # مسح صفوف فارغة تحتوي فقط على مسافات أو أسطر فارغة داخل الخلايا
    html_text = re.sub(r'<tr>\s*(?:<t[hd][^>]*>\s*(?:&nbsp;|\s)*</t[hd]>\s*)+</tr>', '', html_text, flags=re.IGNORECASE)
    
    return html_text

# ══════════════════════════════════════════════════════════
# 🔧 تحويل اتجاه الجداول إلى LTR قبل تصدير الوورد
#    (لأن bidiVisual في DOCX يتولى عكسها للعربية تلقائياً)
# ══════════════════════════════════════════════════════════
def force_tables_ltr_for_export(html_text):
    html_text = re.sub(r'(<table[^>]*?)\bdir\s*=\s*["\']rtl["\']', r'\1dir="ltr"', html_text, flags=re.IGNORECASE)
    return html_text

# ══════════════════════════════════════════════════════════
# 🚀 Local LibreOffice Converter
# ══════════════════════════════════════════════════════════
def local_libreoffice_convert(file_bytes, input_ext, output_ext):
    logger.info(f"🖥️ Local LibreOffice: Converting {input_ext.upper()} to {output_ext.upper()}...")
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = os.path.join(temp_dir, f"input.{input_ext}")
            with open(input_path, 'wb') as f:
                f.write(file_bytes)
            
            profile_dir = os.path.join(temp_dir, "lo_profile")
            
            filters = {
                "docx": "docx:MS Word 2007 XML",
                "xlsx": "xlsx:Calc MS Excel 2007 XML",
                "pptx": "pptx:Impress MS PowerPoint 2007 XML",
                "pdf": "pdf"
            }
            lo_filter = filters.get(output_ext, output_ext)
            
            command = [
                'libreoffice',
                f'-env:UserInstallation=file://{profile_dir}',
                '--headless',
                '--invisible',
                '--nocrashreport',
                '--nodefault',
                '--nofirststartwizard',
                '--nologo',
                '--norestore',
                '--convert-to', lo_filter,
                '--outdir', temp_dir,
                input_path
            ]
            
            env = os.environ.copy()
            env["SAL_USE_VCLPLUGIN"] = "gen"
            env["LC_ALL"] = "C.UTF-8"
            env["DISPLAY"] = ":99"
            
            process = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=120, env=env)
            output_path = os.path.join(temp_dir, f"input.{output_ext}")
            
            if process.returncode == 0 and os.path.exists(output_path):
                with open(output_path, 'rb') as f:
                    result_bytes = f.read()
                logger.info("✅ Local LibreOffice: Conversion successful!")
                return result_bytes, None
            else:
                error_msg = process.stderr.decode('utf-8', errors='ignore').strip()
                if not error_msg:
                    error_msg = process.stdout.decode('utf-8', errors='ignore').strip() or "Unknown error"
                logger.error(f"❌ LibreOffice Failed! Code: {process.returncode}")
                logger.error(f"❌ Error Details: {error_msg}")
                return None, f"خطأ المحرك (Code {process.returncode}): {error_msg}"
    except Exception as e:
        logger.error(f"❌ Local LibreOffice Exception: {str(e)}")
        return None, f"استثناء المحرك: {str(e)}"

# 💡 الرادار اللغوي الذكي (يحدد هل النص عربي أم لاتيني)
def has_arabic(text):
    return bool(re.search(r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]', text))

def get_style_prompt(style, mode):
    global_rules = """
⚠️ DRAFTING VS. FORMATTING & ZERO HALLUCINATION (CRITICAL RULE):
- DRAFTING MODE: If the user asks you to "write", "compose", or "draft" a document based on a brief topic, act as a professional copywriter to structure the letter/document. HOWEVER, YOU MUST STRICTLY USE ONLY THE INFORMATION PROVIDED BY THE USER. DO NOT invent or hallucinate fake names, fake phone numbers, fake prices, or fake company names.
- 🚫 NO PLACEHOLDERS: If the user does not provide a specific piece of information, DO NOT create empty placeholders UNLESS the user explicitly requests a fillable form or blank template. Otherwise, simply OMIT that element entirely.
- FORMATTING MODE: If the user provides ready-made text, a draft, or specific data, your ONLY job is to format their EXACT text into professional HTML. You MUST NOT add, modify, or remove a single word of their content. ZERO hallucination.

⚠️ EXCLUSION RULE:
- 🚫 You MUST completely IGNORE, DELETE, and EXCLUDE any letterheads (headers at the top), footers (at the bottom), logos, stamps, and signatures from the user's uploaded images. DO NOT CREATE FAKE LETTERHEADS.

⚠️ SMART HTML STRUCTURE, DYNAMIC ALIGNMENT & TABLE USAGE (ALL LANGUAGES):
1. 🚫 BREAK THE WALL OF TEXT! A document where every single line starts from the same edge is ugly and rejected. You MUST vary the alignment using inline CSS to make it look like a real printed professional document in ANY language:
   - MAIN TITLES: MUST be strictly centered using `<h1 style="text-align: center;">`.
   - RECIPIENT BLOCK (المرسل إليه / À Monsieur): 
     * For Latin/French/English (e.g., "À monsieur..."): It MUST be aligned to the EXTREME LEFT (`text-align: left; margin-left: 0;`) OR explicitly CENTERED. NEVER push it to the right.
     * For Arabic (e.g., "إلى السيد..."): It MUST be strictly aligned to the EXTREME RIGHT (`text-align: right; margin-right: 0;`). NEVER push it to the left or center it.
   - LONG PARAGRAPHS: MUST be justified to fill the line evenly using `text-align: justify;`.
   - METADATA & FILLABLE FIELDS: Place them intelligently (e.g., using flexbox). If a field is meant to be filled by hand after printing (e.g., "التاريخ:" or "التوقيع:"), you MUST leave sufficient physical writing space next to it (e.g., using "التاريخ: ....................") and NEVER push it to the absolute edge of the page.
   - SIGNATURES / SENDER INFO: Move them to the opposite bottom corner or center them. Do not stack everything on one side.
   - Make the layout dynamic, breathing, and visually balanced!
2. AVOID UNNECESSARY TABLES: DO NOT use tables for standard paragraphs, headers, dates, or signatures. Use tables ONLY for actual tabular data grids (e.g., items, prices, schedules).
3. NO DIV TABLES: When you DO use tables, use classical HTML `<table>`, `<tr>`, `<td>`, `<th>`.
4. 🚫 NO GHOST BOXES: NEVER use CSS `border`, `outline`, or `background` on `<div>`, `<p>`, or `<span>`. Borders are STRICTLY allowed ONLY on `<table>`, `<th>`, and `<td>`.
5. 🚫 NO EMPTY ROWS: NEVER create empty `<tr>` rows or spacer rows at the top of the table. Start directly with the actual text headers. Do NOT use `<thead>`, `<tbody>`, or `<tfoot>` tags.
6. 📊 INVOICE TOTALS (COLSPAN): For rows calculating "Total" (الإجمالي), use the `colspan` attribute to merge empty cells nicely.
7. 🚫 NO FIXED FONT SIZES (CRITICAL FOR DYNAMIC SCALING): NEVER use hardcoded pixel or point sizes (e.g., `font-size: 14px;` or `12pt`). You MUST use standard semantic tags (`<h1>`, `<h2>`) or relative sizes (e.g., `font-size: 1.2em;`, `120%`) for visual hierarchy. The parent container controls the base size.

⚠️ BIDI & MULTILINGUAL LAYOUT LOCKS (MANDATORY):
- Outermost wrapper MUST use `dir="ltr"`.
- Arabic text and `<table>` elements MUST explicitly use `dir="rtl"`.
- Latin/French/English text and `<table>` elements MUST explicitly use `dir="ltr"`.
- 🔄 TABLE COLUMN ORDER: Extract columns in their exact natural logical order as they appear. DO NOT attempt to manually reverse or flip the columns. The browser handles RTL/LTR table rendering automatically based on the `dir` attribute.
- NUMBER ANTI-REVERSAL: ALL numbers MUST strictly be wrapped in: `<span dir="ltr" style="display:inline-block; direction:ltr; unicode-bidi:isolate; white-space:nowrap;"></span>`.
"""
    if mode == "simulation":
        return f"""CLONING: Reproduce EXACTLY text/tables from the reference image.
IGNORE logos, stamps, signatures. Do NOT invent data.
⚠️ EXCEPTIONAL SCENARIO: If the image is a SINGLE circular stamp, produce ONLY an inline <svg> element.
{global_rules}
RULE E – NO BORDERS: You MUST NOT add any outer border, stroke, or page-like box.
RULE F – CAMERA DISTORTION: Ignore physical distortion. Reconstruct in its NATURAL format adapting to the canvas."""

    design_base = ""
    if style == "modern":
        design_base = """MODERN/CREATIVE - Professional, beautiful, and highly aesthetic document design.
CREATIVE FREEDOM: Choose harmonious modern color palettes, elegant typography. Use soft background colors for table headers."""
    else:
        design_base = """FORMAL/OFFICIAL - Ultra clean, strictly official document design.
⚠️ CRITICAL HEADINGS RULE: ABSOLUTELY NO vertical lines, NO border-left, NO border-right, and NO blockquotes next to any headings. Headings MUST be plain, clean, bold text.
⚠️ CRITICAL TABLE RULE: STRICTLY use plain `<table>` with pure black borders. NO background colors, NO gray cells, NO shaded rows. Keep it 100% formal, printable, and transparent.
TYPOGRAPHY: Dynamic sizes. Title bold centered."""

    return f"{design_base}\n\n{global_rules}"


def detect_document_type(user_msg):
    msg_lower = user_msg.lower()
    single_page_keywords = ['فاتورة', 'facture', 'invoice', 'devis', 'عرض سعر', 'bon', 'شهادة', 'certificate', 'attestation', 'رسالة', 'letter', 'lettre', 'courrier', 'إيصال', 'receipt', 'reçu', 'تصريح', 'declaration', 'إذن', 'autorisation', 'بطاقة', 'card']
    multi_page_keywords = ['تقرير', 'report', 'rapport', 'دراسة', 'study', 'étude', 'بحث', 'research', 'خطة', 'plan', 'مشروع', 'project', 'تفصيلي', 'detailed', 'شامل', 'comprehensive']
    for kw in single_page_keywords:
        if kw in msg_lower: return "single_page"
    for kw in multi_page_keywords:
        if kw in msg_lower: return "multi_page"
    return "auto"


@app.route("/", methods=["GET"])
def index():
    return jsonify({"status": "Monjez V10 Server Active", "features": ["documents", "simulation", "design", "translation", "word_export", "magic_convert"]})

@app.route("/gemini", methods=["POST"])
def generate():
    if not get_client(): return jsonify({"error": "Gemini API Offline"}), 500
    try:
        data = request.json
        user_msg = data.get("message", "")
        mode = data.get("mode", "documents")
        style = data.get("style", "formal")
        page_size = data.get("pageSize", "a4Portrait")
        reference_b64 = data.get("reference_image")
        letterhead_b64 = data.get("letterhead_image")

        style_prompt = get_style_prompt(style, mode)
        doc_type = detect_document_type(user_msg)

        page_dimensions = {
            "a4Portrait": {"w": 595, "h": 842, "orientation": "portrait", "physical": "21.0cm x 29.7cm"},
            "a4Landscape": {"w": 842, "h": 595, "orientation": "landscape", "physical": "29.7cm x 21.0cm"},
            "a3": {"w": 842, "h": 1191, "orientation": "portrait A3", "physical": "29.7cm x 42.0cm"},
            "a5": {"w": 420, "h": 595, "orientation": "portrait A5", "physical": "14.8cm x 21.0cm"},
        }
        page_info = page_dimensions.get(page_size, page_dimensions["a4Portrait"])
        is_landscape = page_info["w"] > page_info["h"]

        landscape_extra = f" LANDSCAPE LAYOUT: Tables MUST fit within this width horizontally, but text can flow naturally downwards." if is_landscape else ""
        orientation_instruction = f"PAGE FORMAT: {page_info['orientation']} — Physical Canvas Size: {page_info['physical']} (Target width: {page_info['w']}px). {landscape_extra}"
        
        ref_note = "\nATTACHED IMAGE: Insert using <img src='data:image/jpeg;base64,...' style='max-width:80%; height:auto; margin:8px auto; display:block;' />" if reference_b64 and mode != "simulation" else ""

        doc_type_instruction = "SINGLE-PAGE DOCUMENT: Optimize space beautifully on one page." if doc_type == "single_page" else "MULTI-PAGE DOCUMENT: Allow natural flow across multiple pages."

        svg_rule = "NO `<html>`, `<body>`. (EXCEPTION: `<svg>` is ONLY allowed for the standalone circular stamp scenario)." if mode == "simulation" else "NO `<svg>`, `<html>`, `<body>`."

        prompt = f"""You are a STRICT Document Formatter.
{style_prompt}
{orientation_instruction}
{ref_note}
{doc_type_instruction}
TECHNICAL RULES:
1. PURE HTML ONLY. Just `<div>`, `<table>`, `<h1>`, `<p>`. {svg_rule}
2. NO BORDERS AROUND DOCUMENT.
3. WRAPPER CONFIG: The outermost wrapper MUST NOT have excessive padding. Use `<div style="width:100%; max-width:100%; margin:0 auto; padding:5px; box-sizing:border-box; direction:ltr; overflow-wrap:anywhere; word-break:break-word; overflow:hidden;">`.
OUTPUT: Return raw HTML only."""

        contents = [user_msg] if user_msg else ["Create a formal document."]
        if reference_b64:
            contents.append(get_types().Part.from_bytes(data=base64.b64decode(reference_b64), mime_type="image/jpeg"))
        if letterhead_b64:
            contents.append("Ensure layout fits empty space below this letterhead.")
            contents.append(get_types().Part.from_bytes(data=base64.b64decode(letterhead_b64), mime_type="image/jpeg"))

        gen_config = get_types().GenerateContentConfig(system_instruction=prompt, temperature=0.15, max_output_tokens=20000)

        try:
            resp = call_gemini("gemini-3-flash-preview", contents, gen_config, 55)
        except:
            resp = call_gemini("gemini-2.5-flash", contents, gen_config, 50)

        clean_html = clean_html_output(resp.text or "")
        used_tokens = extract_tokens(resp)
        logger.info(f"✅ Generated HTML (mode: {mode}, page: {page_size}) | Tokens: {used_tokens}")
        return jsonify({"response": clean_html, "used_tokens": used_tokens})
    except Exception as e:
        logger.error(f"Error: {str(e)}", exc_info=True)
        return jsonify({"error": "Failed", "details": str(e), "used_tokens": 0}), 500


@app.route("/modify", methods=["POST"])
def modify():
    if not get_client(): return jsonify({"error": "Gemini API Offline"}), 500
    try:
        data = request.json
        current_html = data.get("current_html") or data.get("currentSVG") or data.get("current_svg") or data.get("htmlContent") or ""
        instruction = data.get("instruction") or data.get("prompt") or ""
        ref_b64 = data.get("reference_image")

        if not current_html.strip():
            logger.error("❌ ERROR: current_html is empty!")
            return jsonify({"error": "Failed", "details": "لم يتم العثور على محتوى المستند الحالي لإجراء التعديل الذكي. يرجى المحاولة مرة أخرى."}), 400

        img_note = f"\nINSERT image: <img src='data:image/jpeg;base64,{ref_b64}' style='max-width:80%; height:auto; margin:8px auto; display:block;' />" if ref_b64 else ""

        # 👇 التعديل هنا: دمج القيود الصارمة للإنشاء (الخطوط، الهوامش، المسافات، ومنع الخلفيات)
        sys = f"""You are a STRICT HTML PATCHING ENGINE. You are NOT a designer.
You will receive a <CURRENT_HTML> document and a <USER_REQUEST>.

CRITICAL RULES (MUST FOLLOW STRICTLY):
1. EXACT COPY-PASTE: Output the EXACT SAME HTML structure provided. DO NOT delete unrelated text or sections.
2. SURGICAL EDIT: Apply the exact surgical change requested. DO NOT hallucinate or add fake elements.
3. BIDI & TYPOGRAPHY PROTECTION: 
   - Preserve `dir="ltr"` on wrappers. Arabic `<table>` elements use `dir="rtl"`.
   - Protect phone numbers with `<span dir="ltr" style="display:inline-block; unicode-bidi:bidi-override; white-space:nowrap;">`.
   - Text in Arabic MUST use `font-family: 'Arial', sans-serif;`. Text in Latin/English MUST use `font-family: 'Times New Roman', serif;`.
4. 🚫 NO BORDERS & NO BACKGROUNDS (CRITICAL): NEVER add outer borders, strokes, shadow boxes, or background colors (especially dark ones) to the main wrappers (`<div>`, `<p>`, `<span>`). The document MUST remain a clean, borderless, transparent standard paper layout.
5. NO FORCED SPACING: DO NOT inject inline `line-height` or custom `margin/padding` into text elements (`<p>`, `<span>`) unless the user explicitly asks for it. Rely on the document's global layout.
6. RETURN FULL HTML: Return the complete patched HTML. Do not truncate or use placeholders like ''.
{img_note}

OUTPUT FORMAT:
[MESSAGE]
وصف قصير للتعديل باللغة العربية
[/MESSAGE]
[HTML]
(ضع هنا كود الـ HTML المعدل كاملاً)
[/HTML]"""

        cfg = get_types().GenerateContentConfig(
            system_instruction=sys, 
            temperature=0.0, 
            max_output_tokens=16384
        )

        cts = [f"<CURRENT_HTML>\n{current_html}\n</CURRENT_HTML>\n\n<USER_REQUEST>\n{instruction}\n</USER_REQUEST>\n\nTASK: Apply the exact surgical change and return FULL updated HTML."]
        if ref_b64:
            cts.append(get_types().Part.from_bytes(data=base64.b64decode(ref_b64), mime_type="image/jpeg"))

        try:
            resp = call_gemini("gemini-3-flash-preview", cts, cfg, 55)
        except:
            resp = call_gemini("gemini-2.5-flash", cts, cfg, 50)

        used_tokens = extract_tokens(resp)
        text = resp.text or ""
        msg_match = re.search(r'\[MESSAGE\](.*?)\[/MESSAGE\]', text, re.DOTALL | re.IGNORECASE)
        html_match = re.search(r'\[HTML\](.*?)\[/HTML\]', text, re.DOTALL | re.IGNORECASE)

        if html_match:
            new_inner = html_match.group(1).strip()
            message = msg_match.group(1).strip() if msg_match else "تم التعديل بنجاح ✨"
        else:
            new_inner = clean_html_output(text)
            new_inner = re.sub(r'\[MESSAGE\].*?\[/MESSAGE\]', '', new_inner, flags=re.DOTALL | re.IGNORECASE).strip()
            message = "تم التعديل بنجاح ✨"

        return jsonify({"response": new_inner, "message": message, "used_tokens": used_tokens})
    except Exception as e:
        logger.error(f"Modify Error: {str(e)}", exc_info=True)
        return jsonify({"error": "Failed", "details": str(e), "used_tokens": 0}), 500


@app.route("/format", methods=["POST"])
def smart_format():
    if not get_client(): return jsonify({"error": "Gemini API Offline"}), 500
    try:
        data = request.json
        current_html = data.get("current_html", "")

        sys = f"""You are a STRICT Document Editor. The user has manually edited this document.
YOUR MISSION:
1. CLEANUP & STRUCTURE: Wrap loose text in proper tags. Apply logical Alignments.
2. STRICT PRESERVATION: NEVER delete, alter, or add to the actual facts, text, or meaning. NO HALLUCINATION.
3. BIDI FIX: Ensure wrappers use `dir="ltr"`. Arabic `<table>` elements use `dir="rtl"`. Arabic text uses `dir="rtl" style="text-align:right"`. Protect phone numbers.
OUTPUT FORMAT:
[MESSAGE]
تم تنسيق وترتيب المستند بنجاح ✨
[/MESSAGE]
[HTML]
(ضع هنا كود الـ HTML المنسق كاملاً)
[/HTML]"""

        cfg = get_types().GenerateContentConfig(system_instruction=sys, temperature=0.0, max_output_tokens=16384)
        cts = [f"<MESSY_HTML>\n{current_html}\n</MESSY_HTML>\n\nPlease format and fix Bidi issues professionally without changing text."]

        try:
            resp = call_gemini("gemini-3-flash-preview", cts, cfg, 55)
        except:
            resp = call_gemini("gemini-2.5-flash", cts, cfg, 50)

        used_tokens = extract_tokens(resp)
        text = resp.text or ""
        msg_match = re.search(r'\[MESSAGE\](.*?)\[/MESSAGE\]', text, re.DOTALL | re.IGNORECASE)
        html_match = re.search(r'\[HTML\](.*?)\[/HTML\]', text, re.DOTALL | re.IGNORECASE)

        if html_match:
            new_inner = html_match.group(1).strip()
            message = msg_match.group(1).strip() if msg_match else "تم التنسيق ✨"
        else:
            new_inner = clean_html_output(text)
            new_inner = re.sub(r'\[MESSAGE\].*?\[/MESSAGE\]', '', new_inner, flags=re.DOTALL | re.IGNORECASE).strip()
            message = "تم التنسيق ✨"

        logger.info("✅ Document Smartly Formatted")
        return jsonify({"response": new_inner, "message": message, "used_tokens": used_tokens})
    except Exception as e:
        logger.error(f"Format Error: {str(e)}", exc_info=True)
        return jsonify({"error": "Failed", "details": str(e), "used_tokens": 0}), 500


# ══════════════════════════════════════════════════════════
# مسار تحويل HTML/PDF إلى Word 
# ══════════════════════════════════════════════════════════
@app.route("/convert_to_word", methods=["POST"])
def convert_to_word():
    try:
        data = request.json
        html_content = data.get("html_content") or data.get("current_html")
        pdf_b64 = data.get("pdf_base64", "")
        letterhead_b64 = data.get("letterhead_base64", "") 
        letterhead_on_all_pages = data.get("letterhead_on_all_pages", False)
        
        input_fmt = "html"
        used_tokens = 0

        if pdf_b64 and not html_content:
            logger.info("📄 Converting PDF to Word via AI Bridge first (To preserve tables)...")
            gemini_bytes = base64.b64decode(pdf_b64)
            
            bridge_prompt = """You are an OCR and Document Extraction Engine.
Your task is to precisely extract ALL content from the attached document and convert it into a fully structured, professional HTML document.
CRITICAL RULES:
1. NO HALLUCINATIONS: Extract the exact words, numbers, and tables. Do not summarize or invent text.
2. 🚫 CRITICAL EXCLUSION RULE: IGNORE, DELETE, and EXCLUDE any letterheads, footers, logos, stamps, and signatures.
3. TABLES & COLSPAN: Use proper `<table>`. NO background colors. For "Total" (الإجمالي) rows, use `colspan` nicely.
4. 🚫 NO EMPTY ROWS: NEVER create empty `<tr>` rows or spacer cells. Do NOT use `<thead>`, `<tbody>`, or `<tfoot>` tags.
5. 🚫 NO GHOST BOXES: NEVER use CSS borders on `<div>`, `<p>`, or `<span>`.
6. 🔄 COLUMN ORDER: Extract columns exactly as they appear in their natural logical order without reversing them.
7. NUMBERS: Wrap any standalone numbers/dates in `<span dir="ltr"></span>`.
8. NO MARKDOWN: Output strictly pure HTML code."""
            
            contents = [bridge_prompt, get_types().Part.from_bytes(data=gemini_bytes, mime_type="application/pdf")]
            gen_config = get_types().GenerateContentConfig(temperature=0.0, max_output_tokens=16384)
            
            try: resp = call_gemini("gemini-3-flash-preview", contents, gen_config, 90)
            except: resp = call_gemini("gemini-2.5-flash", contents, gen_config, 90)
            
            used_tokens = extract_tokens(resp)
            extracted_html = clean_html_output(resp.text or "")
            if not extracted_html:
                return jsonify({"error": "Failed", "details": "فشل الذكاء الاصطناعي في قراءة الـ PDF.", "used_tokens": used_tokens}), 500
            
            html_content = extracted_html

        if html_content:
            logger.info("📄 Preparing HTML for LibreOffice Word Conversion...")

            html_content = force_table_borders(html_content)
            html_content = force_tables_ltr_for_export(html_content)
            html_content = re.sub(r'font-family\s*:[^;"]+[;]?', '', html_content, flags=re.IGNORECASE)
            
            # 💡 لحام الأرقام لمنع انعكاسها
            html_content = re.sub(r'(\d)\s+(?=\d)', r'\1&nbsp;', html_content)
            
            is_arabic_doc = has_arabic(html_content)
            body_dir = "rtl" if is_arabic_doc else "ltr"

            html_content = re.sub(
                r'<div[^>]*display\s*:\s*flex[^>]*>.*?<div[^>]*border-bottom[^>]*>.*?</div>.*?<div[^>]*>\s*:\s*</div>.*?<div[^>]*>(.*?)</div>.*?</div>',
                r'<p dir="rtl" style="text-align:right; margin:0;">\1: ........................................</p>',
                html_content, flags=re.IGNORECASE | re.DOTALL)
            html_content = re.sub(
                r'<div[^>]*display\s*:\s*flex[^>]*>.*?<div[^>]*>(.*?)</div>.*?<div[^>]*>\s*:\s*</div>.*?<div[^>]*border-bottom[^>]*>.*?</div>.*?</div>',
                r'<p dir="rtl" style="text-align:right; margin:0;">\1: ........................................</p>',
                html_content, flags=re.IGNORECASE | re.DOTALL)
            html_content = re.sub(r'<div[^>]*border-bottom[^>]*>(\s|&nbsp;)*</div>', ' ........................................ ', html_content, flags=re.IGNORECASE)

            full_html = f"""<html lang="ar" dir="{body_dir}">
<head>
<meta charset="utf-8">
<style>
  * {{ font-family: 'Arial', sans-serif !important; }}
  table {{ border-collapse: collapse; margin: 10px 0; width: 100% !important; }}
  th, td {{ border: 1px solid #000000; padding: 4px !important; line-height: 1.1 !important; vertical-align: middle !important; }}
  p, h1, h2, h3, h4, h5, h6, div, span {{ margin: 0; padding: 0; border: none !important; background: transparent !important; line-height: 1.2 !important; }}
</style>
</head>
<body>
{html_content}
</body>
</html>"""
            file_bytes = full_html.encode('utf-8')
            
        else:
            return jsonify({"error": "Failed", "details": "لم يتم إرسال محتوى المستند.", "used_tokens": 0}), 400

        raw_docx_bytes, err_msg = local_libreoffice_convert(file_bytes, input_fmt, "docx")
        
        if not raw_docx_bytes:
            return jsonify({"error": "Failed", "details": f"فشل LibreOffice: {err_msg}", "used_tokens": used_tokens}), 500

        # ══════════════════════════════════════════════════════════
        # معالجة الوورد: الحرية للمحاذاة وإصلاح انعكاس الأعمدة والتاريخ
        # ══════════════════════════════════════════════════════════
        logger.info("💉 Local Processing: Deep XML Fixes for Fonts and Tables...")
        doc_stream = io.BytesIO(raw_docx_bytes)
        doc = docx.Document(doc_stream)
        section = doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Cm(4.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

        def clean_and_format_paragraph(paragraph, is_table=False):
            text = paragraph.text.strip()
            is_arabic_para = has_arabic(text) if text else is_arabic_doc

            pPr = paragraph._element.get_or_add_pPr()
            
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = OxmlElement('w:spacing')
                pPr.append(spacing)
            
            # ✅ ضبط التباعد (مرونة للنص العادي، وتماسك للجداول)
            if is_table:
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '60')  # تباعد خفيف للجدول
                spacing.set(qn('w:line'), '280')  
            else:
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '240') # تباعد واضح بين الفقرات (12pt)
                spacing.set(qn('w:line'), '360')  # تباعد أسطر يعادل 1.5 للنصوص العادية
            spacing.set(qn('w:lineRule'), 'auto')

            # ✅ ضبط حجم الخط (30 تعني 15pt للنص العادي | 24 تعني 12pt للجدول)
            target_size = '24' if is_table else '30'

            for run in paragraph.runs:
                rPr = run._element.get_or_add_rPr()
                
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:cs'), 'Arial')
                rFonts.set(qn('w:ascii'), 'Arial')
                rFonts.set(qn('w:hAnsi'), 'Arial')
                rFonts.set(qn('w:eastAsia'), 'Arial')
                
                sz = rPr.find(qn('w:sz'))
                if sz is None: sz = OxmlElement('w:sz'); rPr.append(sz)
                sz.set(qn('w:val'), target_size)
                
                szCs = rPr.find(qn('w:szCs'))
                if szCs is None: szCs = OxmlElement('w:szCs'); rPr.append(szCs)
                szCs.set(qn('w:val'), target_size)

        for table in doc.tables:
            table.autofit = True
            tblPr = table._element.tblPr
            if tblPr is not None:
                tblLayout = tblPr.find(qn('w:tblLayout'))
                if tblLayout is not None:
                    tblLayout.set(qn('w:type'), 'autofit')
                
                tblW = tblPr.find(qn('w:tblW'))
                if tblW is None:
                    tblW = OxmlElement('w:tblW')
                    tblPr.append(tblW)
                tblW.set(qn('w:w'), '5000') 
                tblW.set(qn('w:type'), 'pct')

                # 💡 [الحل الجذري المعتمد]: إجبار الجدول على الاتجاه العربي في الوورد ليمنع انعكاس الأعمدة
                if is_arabic_doc:
                    bidiVisual = tblPr.find(qn('w:bidiVisual'))
                    if bidiVisual is None:
                        bidiVisual = OxmlElement('w:bidiVisual')
                        tblPr.append(bidiVisual)

            for row in table.rows:
                trPr = row._tr.get_or_add_trPr()
                for trHeight in trPr.findall(qn('w:trHeight')):
                    trPr.remove(trHeight)
                
                for cell in row.cells:
                    tcPr = cell._element.get_or_add_tcPr()
                    
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is not None:
                        tcW.set(qn('w:type'), 'auto')

                    tcMar = tcPr.find(qn('w:tcMar'))
                    if tcMar is not None: tcPr.remove(tcMar)
                    tcMar = OxmlElement('w:tcMar')
                    for attr in ['top', 'bottom']:
                        node = OxmlElement(f'w:{attr}')
                        node.set(qn('w:w'), '0')
                        node.set(qn('w:type'), 'dxa')
                        tcMar.append(node)
                    tcPr.append(tcMar)

                    for p in cell.paragraphs:
                        clean_and_format_paragraph(p, is_table=True)

        for p in doc.paragraphs:
            clean_and_format_paragraph(p, is_table=False)

        if letterhead_b64:
            header_img_data = base64.b64decode(letterhead_b64)
            header_img_stream = io.BytesIO(header_img_data)
            
            if not letterhead_on_all_pages:
                section.different_first_page_header_footer = True
                target_header = section.first_page_header
            else:
                target_header = section.header
                
            if not target_header.paragraphs:
                target_header.add_paragraph()
            paragraph = target_header.paragraphs[0]
            
            run = paragraph.add_run()
            shape = run.add_picture(header_img_stream, width=Inches(8.27), height=Inches(11.69))
            
            inline = shape._inline
            anchor = OxmlElement('wp:anchor')
            anchor.set('distT', '0'); anchor.set('distB', '0'); anchor.set('distL', '0'); anchor.set('distR', '0')
            anchor.set('simplePos', '0'); anchor.set('relativeHeight', '0'); anchor.set('behindDoc', '1') 
            anchor.set('locked', '0'); anchor.set('layoutInCell', '1'); anchor.set('allowOverlap', '1')

            simplePos = OxmlElement('wp:simplePos')
            simplePos.set('x', '0'); simplePos.set('y', '0')
            anchor.append(simplePos)

            positionH = OxmlElement('wp:positionH')
            positionH.set('relativeFrom', 'page')
            alignH = OxmlElement('wp:align'); alignH.text = 'center'
            positionH.append(alignH); anchor.append(positionH)
            
            positionV = OxmlElement('wp:positionV')
            positionV.set('relativeFrom', 'page')
            alignV = OxmlElement('wp:align'); alignV.text = 'top'
            positionV.append(alignV); anchor.append(positionV)
            
            anchor.append(inline.extent)
            effectExtent = OxmlElement('wp:effectExtent')
            effectExtent.set('l', '0'); effectExtent.set('t', '0'); effectExtent.set('r', '0'); effectExtent.set('b', '0')
            anchor.append(effectExtent)
            
            anchor.append(OxmlElement('wp:wrapNone'))
            anchor.append(inline.docPr)
            anchor.append(OxmlElement('wp:cNvGraphicFramePr'))
            anchor.append(inline.graphic)
            
            inline.getparent().replace(inline, anchor)

        final_docx_stream = io.BytesIO()
        doc.save(final_docx_stream)
        docx_bytes = final_docx_stream.getvalue()
        docx_b64 = base64.b64encode(docx_bytes).decode('utf-8')

        logger.info(f"✅ Final Word Document generated successfully ({len(docx_bytes)} bytes)")
        return jsonify({"docx_base64": docx_b64, "message": "تم التحويل إلى Word بنجاح ✨", "used_tokens": used_tokens})
    except Exception as e:
        logger.error(f"Word Error: {str(e)}", exc_info=True)
        return jsonify({"error": "Failed", "details": f"فشل التحويل: {str(e)}", "used_tokens": 0}), 500


# ══════════════════════════════════════════════════════════
# مسار MAGIC CONVERTER (المحول الشامل)
# ══════════════════════════════════════════════════════════
@app.route("/magic_convert", methods=["POST"])
def magic_convert():
    try:
        data = request.json
        file_b64 = data.get("fileBase64")
        mime_type = data.get("mimeType", "")
        target_format = data.get("targetFormat", "word")
        is_arabic = data.get("isArabic", True)
        extract_only = data.get("extractOnly", False)  

        if not file_b64:
            return jsonify({"error": "Failed", "details": "لم يتم العثور على الملف", "used_tokens": 0}), 400

        mime_lower = mime_type.lower()
        input_ext = "pdf"
        if "html" in mime_lower: input_ext = "html"
        elif "jpeg" in mime_lower or "jpg" in mime_lower: input_ext = "jpg"
        elif "png" in mime_lower: input_ext = "png"
        elif "msword" in mime_lower or "word" in mime_lower or "docx" in mime_lower: input_ext = "docx" 
        elif "excel" in mime_lower or "xls" in mime_lower or "spreadsheet" in mime_lower: input_ext = "xlsx"
        elif "powerpoint" in mime_lower or "ppt" in mime_lower or "presentation" in mime_lower: input_ext = "pptx"
        
        file_bytes = base64.b64decode(file_b64)
        
        output_ext = "docx"
        if target_format == "excel": output_ext = "xlsx"
        elif target_format == "powerpoint": output_ext = "pptx"
        elif target_format == "pdf": output_ext = "pdf"
        elif target_format == "html": output_ext = "html"

        logger.info(f"🔄 Magic Request: {input_ext.upper()} ➡️ {output_ext.upper()}")

        direct_conversions = [
            ("docx", "pdf"), ("doc", "pdf"),
            ("xlsx", "pdf"), ("xls", "pdf"),
            ("pptx", "pdf"), ("ppt", "pdf"),
            ("html", "docx"), ("html", "xlsx"), ("html", "pdf")
        ]
        
        used_tokens = 0

        if (input_ext, output_ext) in direct_conversions and not extract_only:
            logger.info("⚡ Route 1: Direct LibreOffice Conversion (No AI needed)...")
            
            if input_ext == "html":
                html_text = file_bytes.decode('utf-8')
                html_text = force_table_borders(html_text)
                html_text = force_tables_ltr_for_export(html_text)
                html_text = re.sub(r'font-family\s*:[^;"]+[;]?', '', html_text, flags=re.IGNORECASE)
                
                html_text = re.sub(r'(\d)\s+(?=\d)', r'\1&nbsp;', html_text)
                is_arabic_doc = has_arabic(html_text)
                body_dir = "rtl" if is_arabic_doc else "ltr"
                
                full_html = f"""<html lang="ar" dir="{body_dir}"><head><meta charset="utf-8">
<style>* {{ font-family: 'Arial', sans-serif !important; }} table {{ border-collapse: collapse; margin: 10px 0; width: 100% !important; }} th, td {{ border: 1px solid #000; padding: 4px !important; line-height: 1.1 !important; }} p, h1, h2, h3, h4, h5, h6, div, span {{ margin: 0; padding: 0; border: none !important; background: transparent !important; }}</style>
</head><body>{html_text}</body></html>"""
                file_bytes = full_html.encode('utf-8')

            result_bytes, err_msg = local_libreoffice_convert(file_bytes, input_ext, output_ext)
            
            if result_bytes:
                result_b64 = base64.b64encode(result_bytes).decode('utf-8')
                return jsonify({
                    "file_base64": result_b64,
                    "extension": output_ext,
                    "message": f"تم التحويل إلى {target_format.upper()} بنجاح ✨",
                    "used_tokens": used_tokens
                })
            else:
                logger.warning(f"⚠️ Direct conversion failed: {err_msg}. Falling back to AI Route if applicable.")

        logger.info("🧠 Route 2: AI OCR & Extraction Bridge...")
        gemini_bytes = file_bytes
        gemini_mime = "application/pdf"
        
        if input_ext in ["docx", "doc", "xlsx", "xls", "pptx", "ppt"]:
            logger.info("🔄 Converting Document to PDF first via LibreOffice for AI Reading...")
            gemini_bytes, err_msg = local_libreoffice_convert(file_bytes, input_ext, "pdf")
            if not gemini_bytes:
                return jsonify({"error": "Failed", "details": f"فشل تجهيز المستند للقراءة: {err_msg}", "used_tokens": 0}), 500
            gemini_mime = "application/pdf"
        elif input_ext in ["jpg", "jpeg"]: gemini_mime = "image/jpeg"
        elif input_ext == "png": gemini_mime = "image/png"

        target_focus = "tables and grids format specifically for Excel" if output_ext == "xlsx" else "general document structure"
        
        bridge_prompt = f"""You are an elite OCR and Document Extraction Engine.
Your task is to precisely extract ALL content from the attached document and convert it into a fully structured, professional HTML document. Focus on {target_focus}.

CRITICAL RULES:
1. NO HALLUCINATIONS: Extract the exact words, numbers, and tables. Do not summarize or invent text.
2. 🚫 CRITICAL EXCLUSION RULE: IGNORE, DELETE, and EXCLUDE any letterheads, footers, logos, stamps, and signatures.
3. TABLES FOR GRIDS ONLY: Use `<table>` ONLY for actual tabular data (items, prices, schedules). Regular text, headers, and dates MUST be in `<p>` or `<div>`. NEVER put the whole document in a table.
4. COLSPAN: For "Total" (الإجمالي) rows, use `colspan` elegantly.
5. 🚫 NO EMPTY ROWS: NEVER create empty `<tr>` or `<th>` rows. Start directly with the text headers. Do NOT use `<thead>`, `<tbody>`, or `<tfoot>` tags.
6. 🚫 NO GHOST BOXES: NEVER use CSS borders on `<div>`, `<p>`, or `<span>`. Borders are for tables ONLY.
7. 🔄 COLUMN ORDER: Extract columns exactly as they appear in their natural logical order without reversing them.
8. NUMBERS: Wrap standalone numbers/dates in `<span dir="ltr"></span>`.
9. PURE HTML ONLY. Do not wrap in ```html."""
        
        contents = [bridge_prompt, get_types().Part.from_bytes(data=gemini_bytes, mime_type=gemini_mime)]
        gen_config = get_types().GenerateContentConfig(temperature=0.0, max_output_tokens=16384)
        
        try: resp = call_gemini("gemini-3-flash-preview", contents, gen_config, 90)
        except: resp = call_gemini("gemini-2.5-flash", contents, gen_config, 90)
        
        used_tokens = extract_tokens(resp)
        extracted_html = clean_html_output(resp.text or "")
        if not extracted_html:
            return jsonify({"error": "Failed", "details": "فشل الذكاء الاصطناعي في قراءة الملف", "used_tokens": used_tokens}), 500
        
        if extract_only or target_format == "html":
            return jsonify({"html_content": extracted_html, "message": "تم استخراج النصوص بنجاح ✨", "used_tokens": used_tokens})
        
        logger.info(f"📄 Wrapping extracted HTML to final format: {output_ext.upper()}...")
        
        extracted_html = force_table_borders(extracted_html)
        extracted_html = force_tables_ltr_for_export(extracted_html)
        extracted_html = re.sub(r'(\d)\s+(?=\d)', r'\1&nbsp;', extracted_html)
        
        is_arabic_doc = has_arabic(extracted_html)
        body_dir = "rtl" if is_arabic_doc else "ltr"
        
        full_html = f"""<html lang="ar" dir="{body_dir}"><head><meta charset="utf-8">
<style>* {{ font-family: 'Arial', sans-serif !important; }} table {{ border-collapse: collapse; margin: 10px 0; width: 100% !important; }} th, td {{ border: 1px solid #000; padding: 4px !important; line-height: 1.1 !important; }} p, h1, h2, h3, h4, h5, h6, div, span {{ margin: 0; padding: 0; border: none !important; background: transparent !important; line-height: 1.2 !important; }}</style>
</head><body>{extracted_html}</body></html>"""
        
        final_bytes = full_html.encode('utf-8')
        result_bytes, err_msg = local_libreoffice_convert(final_bytes, "html", output_ext)
        
        if not result_bytes:
            return jsonify({"error": "Failed", "details": f"فشل تجميع الملف النهائي: {err_msg}", "used_tokens": used_tokens}), 500
            
        result_b64 = base64.b64encode(result_bytes).decode('utf-8')
        return jsonify({
            "file_base64": result_b64,
            "extension": output_ext,
            "message": f"تم التحويل إلى {target_format.upper()} بنجاح ✨",
            "used_tokens": used_tokens
        })

    except Exception as e:
        logger.error(f"Magic Convert Error: {str(e)}", exc_info=True)
        return jsonify({"error": "Failed", "details": str(e), "used_tokens": 0}), 500


# ══════════════════════════════════════════════════════════
# مسار TRANSLATE DOCUMENT
# ══════════════════════════════════════════════════════════
@app.route("/translate_document", methods=["POST"])
def translate_document():
    if not get_client(): return jsonify({"error": "Gemini API Offline"}), 500
    try:
        data = request.json
        target_language = data.get("target_language", "العربية")
        reference_b64 = data.get("reference_image")
        page_size = data.get("pageSize", "a4Portrait")

        page_dimensions = {
            "a4Portrait": {"w": 595, "h": 842, "orientation": "portrait", "physical": "21.0cm x 29.7cm"},
            "a4Landscape": {"w": 842, "h": 595, "orientation": "landscape", "physical": "29.7cm x 21.0cm"},
            "a3": {"w": 842, "h": 1191, "orientation": "portrait A3", "physical": "29.7cm x 42.0cm"},
            "a5": {"w": 420, "h": 595, "orientation": "portrait A5", "physical": "14.8cm x 21.0cm"},
        }
        page_info = page_dimensions.get(page_size, page_dimensions["a4Portrait"])
        is_landscape = page_info["w"] > page_info["h"]

        landscape_extra = f" LANDSCAPE LAYOUT: Tables MUST fit within this width horizontally, but text can flow naturally downwards." if is_landscape else ""
        orientation_instruction = f"PAGE FORMAT: {page_info['orientation']} — Physical Canvas Size: {page_info['physical']} (Target width: {page_info['w']}px). {landscape_extra}"

        bidi_rules = """
⚠️ BIDI & LAYOUT LOCKS:
- Outermost wrapper MUST use `dir="ltr"`.
- Arabic `<table>` elements MUST use `dir="rtl"`.
- Non-Arabic (Latin/French) `<table>` elements MUST use `dir="ltr"`.
- Arabic text MUST explicitly use `dir="rtl" style="text-align: right;"`
- TABLE COLUMN ORDER: Output HTML columns in their NATURAL logical order exactly as they appear. DO NOT manually reverse the columns.
- NUMBER ANTI-REVERSAL: ALL numbers MUST strictly be wrapped in: `<span dir="ltr" style="display:inline-block; direction:ltr; unicode-bidi:isolate; white-space:nowrap;"></span>`.
"""

        prompt = f"""You are an Expert Professional Translator and Strict Document Formatter.
YOUR MISSION:
1. Clone the exact layout, structure, and tables of the provided document image.
2. TRANSLATE all text into {target_language} with high professional accuracy. 
3. DO NOT invent fake data, logos, or headers. Translate exactly what is there.
4. 🚫 CRITICAL EXCLUSION RULE: You MUST completely IGNORE, DELETE, and EXCLUDE any original letterheads, footers, logos, stamps, and signatures.
{bidi_rules}
{orientation_instruction}
TECHNICAL RULES:
1. PURE HTML ONLY. Just `<div>`, `<table>`, `<h1>`, `<p>`. NO `<svg>`, `<html>`, `<body>`.
2. NO BORDERS AROUND DOCUMENT.
OUTPUT: Return raw HTML only."""

        contents = [f"Translate this document to {target_language} while keeping the exact layout."]
        if reference_b64:
            contents.append(get_types().Part.from_bytes(data=base64.b64decode(reference_b64), mime_type="image/jpeg"))
        else:
            return jsonify({"error": "Failed", "details": "لم يتم إرفاق المستند", "used_tokens": 0}), 400

        gen_config = get_types().GenerateContentConfig(system_instruction=prompt, temperature=0.15, max_output_tokens=20000)

        try:
            resp = call_gemini("gemini-3-flash-preview", contents, gen_config, 55)
        except:
            resp = call_gemini("gemini-2.5-flash", contents, gen_config, 50)

        used_tokens = extract_tokens(resp)
        clean_html = clean_html_output(resp.text or "")
        logger.info(f"✅ Generated Translation HTML (Target: {target_language}) | Tokens: {used_tokens}")
        return jsonify({"response": clean_html, "used_tokens": used_tokens})
    except Exception as e:
        logger.error(f"Error: {str(e)}", exc_info=True)
        return jsonify({"error": "Failed", "details": str(e), "used_tokens": 0}), 500

# ══════════════════════════════════════════════════════════
# 🚀 NEW: DESIGN GENERATION (Vertex AI: Imagen 4 Ultra -> 3 Fallback)
# ══════════════════════════════════════════════════════════

@@app.route("/generate_image", methods=["POST"])
def generate_image():
    import urllib.request
    import urllib.error
    import json
    
    try:
        # ✅ الإبقاء على اسم المفتاح بالشرطة كما هو في بيئتك
        k = os.environ.get("GOOGLE_API-KEY2") or os.environ.get("GOOGLE_API_KEY")
        if not k:
            return jsonify({"error": "Failed", "details": "مفتاح GOOGLE_API-KEY2 غير موجود."}), 500

        data = request.json
        user_prompt = data.get("prompt", "")

        if not user_prompt.strip():
            return jsonify({"error": "Failed", "details": "يرجى كتابة وصف للتصميم المطلوب."}), 400

        logger.info(f"🧠 Step 1: Enhancing prompt via Gemini (Direct REST)...")

        # 🚀 المرحلة 1: المدير الفني الذكي
        # 🚩 تم التوجيه إلى AI Studio ليتوافق مع الـ API Key الخاص بك
        gemini_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={k}"
        sys_instruct = """You are an elite Art Director and Expert Prompt Engineer.
The user will provide a brief idea in Arabic. UNDERSTAND THE CONTEXT and expand it into a MASTERPIECE English prompt for Imagen.
CRITICAL RULES:
1. NO MOCKUPS ALLOWED (STRICTLY FORBIDDEN). DO NOT place designs on walls, paper, screens, 3D objects, or merchandise. Provide the RAW, FLAT, modern, and professional design directly.
2. CONTEXT: IF PRINT (مطبوعات, كرت, فلاير): Clean layout, negative space. IF SOCIAL MEDIA: Visually striking, commercial studio lighting. IF LOGO: Clean, scalable, flat professional design.
3. QUALITY: 8k resolution, cinematic lighting, hyper-realistic photography. NO vector/cartoon unless explicitly requested.
4. CULTURE (STRICT): If people/lifestyle are included, they MUST have authentic Mauritanian facial features and reflect Mauritanian culture (Men MUST wear traditional Daraa/Boubou, Women MUST wear traditional Melhfa). The vibe should be distinctly Mauritanian.
5. OUTPUT ONLY THE ENGLISH PROMPT. No intros."""

        gemini_payload = {
            "contents": [{"role": "user", "parts": [{"text": user_prompt}]}],
            # 🚩 تم تصحيح هيكل systemInstruction ليقبله الخادم
            "systemInstruction": {"parts": [{"text": sys_instruct}]},
            "generationConfig": {"temperature": 0.7}
        }

        try:
            req_gemini = urllib.request.Request(gemini_url, data=json.dumps(gemini_payload).encode('utf-8'), headers={"Content-Type": "application/json"})
            with urllib.request.urlopen(req_gemini, timeout=15) as response:
                gemini_result = json.loads(response.read().decode('utf-8'))
                expanded_prompt = gemini_result["candidates"][0]["content"]["parts"][0]["text"].strip()
                logger.info(f"✨ Super Prompt: {expanded_prompt}")
        except Exception as e:
            logger.warning(f"Gemini enhancement failed, using fallback: {e}")
            expanded_prompt = f"RAW, FLAT design, NO mockups. Ultra-realistic, 8k resolution, Mauritanian cultural context. Subject: {user_prompt}"

        logger.info(f"🎨 Step 2: Generating image using Dynamic Fallback...")

        headers = {"Content-Type": "application/json"}
        payload = {
            "instances": [{"prompt": expanded_prompt}],
            "parameters": {
                "sampleCount": 1,
                "aspectRatio": "1:1",
                "outputOptions": {"mimeType": "image/jpeg"}
            }
        }

        # 🚀 قائمة النماذج من الأقوى إلى الأضمن
        models_to_try = [
            "imagen-4.0-ultra-generate-001",
            "imagen-4.0-generate-001",
            "imagen-3.0-generate-002",
            "imagen-3.0-generate-001"
        ]

        for model_name in models_to_try:
            # 🚩 تم التوجيه إلى AI Studio لتوليد الصور
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:predict?key={k}"
            req = urllib.request.Request(url, data=json.dumps(payload).encode('utf-8'), headers=headers)
            
            try:
                logger.info(f"🚀 Trying {model_name}...")
                with urllib.request.urlopen(req, timeout=45) as response:
                    result = json.loads(response.read().decode('utf-8'))
                    if "predictions" in result and len(result["predictions"]) > 0:
                        img_b64 = result["predictions"][0].get("bytesBase64Encoded")
                        if img_b64:
                            logger.info(f"✅ Design Generated Successfully with {model_name}!")
                            return jsonify({"response": img_b64, "message": "تم التصميم بنجاح ✨"})
            except urllib.error.HTTPError as e:
                logger.warning(f"⚠️ {model_name} unavailable (HTTP {e.code}). Skipping to next...")
                continue 
            except Exception as e:
                logger.warning(f"⚠️ {model_name} failed: {e}. Skipping to next...")
                continue

        # إذا فشلت جميع المحاولات
        logger.error("❌ All image models failed or are unsupported by this key.")
        return jsonify({"error": "Failed", "details": "جميع نماذج الصور غير متاحة حالياً، يرجى المحاولة لاحقاً."}), 500
        
    except Exception as e:
        logger.error(f"Image Gen Error: {str(e)}", exc_info=True)
        return jsonify({"error": "Failed", "details": f"خطأ في الخادم: {str(e)}"}), 500


    except Exception as e:
        logger.error(f"Design Server Error: {str(e)}", exc_info=True)
        return jsonify({"error": "Failed", "details": "حدث خطأ أثناء الاتصال بالخادم."}), 500


# ══════════════════════════════════════════════════════════
# 🌟 مسار ENHANCE TEXT (لتصحيح وتحسين وصف البنود)
# ══════════════════════════════════════════════════════════
@app.route("/enhance_text", methods=["POST"])
def enhance_text():
    if not get_client(): return jsonify({"error": "Gemini API Offline"}), 500
    try:
        data = request.json
        text = data.get("text", "")
        if not text.strip():
            return jsonify({"error": "Failed", "details": "النص فارغ", "used_tokens": 0}), 400

        sys_prompt = """You are an expert corporate billing specialist and strict proofreader.
Analyze the following product/service description from an invoice.

CRITICAL RULES:
1. LANGUAGE: You MUST respond in the EXACT SAME LANGUAGE as the user's input text. Do not translate.
2. TONE: Strictly formal, concise, and professional. NO marketing fluff, NO enthusiasm, NO promotional adjectives (like "amazing", "best", "perfect"). Use precise, standard business/billing terminology suitable for an official corporate invoice.

Return ONLY a valid JSON object with exactly these two keys:
"correction": The original text with only spelling and grammatical errors fixed. Keep the exact original words and meaning.
"suggestion": A refined, highly formal, and concise rewritten version of the text, matching standard corporate billing language.

Do NOT wrap the response in ```json, just return the raw JSON object."""

        cfg = get_types().GenerateContentConfig(
            system_instruction=sys_prompt,
            temperature=0.1,  # تقليل الإبداع لتكون اللغة رسمية وجافة
            response_mime_type="application/json"
        )
        
        contents = [f"Text to enhance: {text}"]
        
        try:
            resp = call_gemini("gemini-3-flash-preview", contents, cfg, 30)
        except:
            resp = call_gemini("gemini-2.5-flash", contents, cfg, 30)
            
        used_tokens = extract_tokens(resp)
        result_text = resp.text.strip()
        
        if result_text.startswith("```json"):
            result_text = result_text[7:-3].strip()
        elif result_text.startswith("```"):
            result_text = result_text[3:-3].strip()
            
        parsed_json = json.loads(result_text)
        parsed_json["used_tokens"] = used_tokens
        return jsonify(parsed_json)
        
    except Exception as e:
        logger.error(f"Enhance Error: {str(e)}", exc_info=True)
        return jsonify({"error": "Failed", "details": str(e), "used_tokens": 0}), 500


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port, threaded=True, debug=False)


# ══════════════════════════════════════════════════════════
# 🌟 مسار ENHANCE TEXT (لتصحيح وتحسين وصف البنود)
# ══════════════════════════════════════════════════════════
@app.route("/enhance_text", methods=["POST"])
def enhance_text():
    if not get_client(): return jsonify({"error": "Gemini API Offline"}), 500
    try:
        data = request.json
        text = data.get("text", "")
        if not text.strip():
            return jsonify({"error": "Failed", "details": "النص فارغ", "used_tokens": 0}), 400

        sys_prompt = """You are an expert corporate billing specialist and strict proofreader.
Analyze the following product/service description from an invoice.

CRITICAL RULES:
1. LANGUAGE: You MUST respond in the EXACT SAME LANGUAGE as the user's input text. Do not translate.
2. TONE: Strictly formal, concise, and professional. NO marketing fluff, NO enthusiasm, NO promotional adjectives (like "amazing", "best", "perfect"). Use precise, standard business/billing terminology suitable for an official corporate invoice.

Return ONLY a valid JSON object with exactly these two keys:
"correction": The original text with only spelling and grammatical errors fixed. Keep the exact original words and meaning.
"suggestion": A refined, highly formal, and concise rewritten version of the text, matching standard corporate billing language.

Do NOT wrap the response in ```json, just return the raw JSON object."""

        cfg = get_types().GenerateContentConfig(
            system_instruction=sys_prompt,
            temperature=0.1,  # تقليل الإبداع لتكون اللغة رسمية وجافة
            response_mime_type="application/json"
        )
        
        contents = [f"Text to enhance: {text}"]
        
        try:
            resp = call_gemini("gemini-3-flash-preview", contents, cfg, 30)
        except:
            resp = call_gemini("gemini-2.5-flash", contents, cfg, 30)
            
        used_tokens = extract_tokens(resp)
        result_text = resp.text.strip()
        
        if result_text.startswith("```json"):
            result_text = result_text[7:-3].strip()
        elif result_text.startswith("```"):
            result_text = result_text[3:-3].strip()
            
        parsed_json = json.loads(result_text)
        parsed_json["used_tokens"] = used_tokens
        return jsonify(parsed_json)
        
    except Exception as e:
        logger.error(f"Enhance Error: {str(e)}", exc_info=True)
        return jsonify({"error": "Failed", "details": str(e), "used_tokens": 0}), 500


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port, threaded=True, debug=False)


